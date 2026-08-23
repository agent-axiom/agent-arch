from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html
from PIL import Image

FIRST_CHAPTERS_IN_PART = {1, 4, 7, 10, 13, 17, 22, 26}
CODE_FONT = "Roboto Mono"
BOOK_TITLE = "Архитектура безопасных ИИ-агентов"
BOOK_SUBJECT = "Проектирование, разработка и эксплуатация безопасных ИИ-агентов"
BOOK_KEYWORDS = "ИИ-агенты, безопасность, архитектура, LLM, AgentOps"
INTERNAL_REFERENCE_RE = re.compile(
    r"\b(?P<label>"
    r"глав(?:а|ы|е|у|ой|ах)|"
    r"рисун(?:ок|ка|ке|ку|ком|ки|ков|кам|ках)|"
    r"таблиц(?:а|ы|е|у|ей|ам|ами|ах)|"
    r"листинг(?:а|е|у|ом|и|ов|ам|ами|ах)?"
    r")\s+(?P<number>[1-9]\d*)\b",
    flags=re.IGNORECASE,
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def markdown_to_html(manuscript: Path, node: Path, node_path: Path) -> str:
    script = r"""
import fs from "node:fs";
import {createRequire} from "node:module";
const require = createRequire(import.meta.url);
const {marked} = await import(require.resolve("marked"));
const source = fs.readFileSync(process.argv[1], "utf8");
process.stdout.write(marked.parse(source, {gfm: true}));
"""
    environment = os.environ.copy()
    environment["NODE_PATH"] = str(node_path)
    result = subprocess.run(
        [str(node), "--input-type=module", "-e", script, str(manuscript)],
        check=True,
        capture_output=True,
        text=True,
        env=environment,
    )
    return result.stdout


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    properties = run._r.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def set_page_break_before(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:pageBreakBefore")) is None:
        properties.append(OxmlElement("w:pageBreakBefore"))


def set_keep_next(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:keepNext")) is None:
        properties.append(OxmlElement("w:keepNext"))


def set_keep_lines(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:keepLines")) is None:
        properties.append(OxmlElement("w:keepLines"))


def begin_complex_field(paragraph, instruction: str) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run.append(instruction_text)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)

    paragraph._p.extend((begin_run, instruction_run, separate_run))


def end_complex_field(paragraph) -> None:
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def add_complex_field(paragraph, instruction: str, placeholder: str) -> None:
    begin_complex_field(paragraph, instruction)

    value_run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = placeholder
    value_run.append(value)

    paragraph._p.append(value_run)
    end_complex_field(paragraph)


def configure_document_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = BOOK_TITLE
    properties.subject = BOOK_SUBJECT
    properties.keywords = BOOK_KEYWORDS
    properties.language = "ru-RU"
    properties.author = ""

    styles = document.styles.element
    defaults = styles.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles.insert(0, defaults)
    run_defaults = defaults.find(qn("w:rPrDefault"))
    if run_defaults is None:
        run_defaults = OxmlElement("w:rPrDefault")
        defaults.append(run_defaults)
    run_properties = run_defaults.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_defaults.append(run_properties)
    language = run_properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run_properties.append(language)
    for attribute in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{attribute}"), "ru-RU")

    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        paragraph_properties = style.element.get_or_add_pPr()
        outline_level = paragraph_properties.find(qn("w:outlineLvl"))
        if outline_level is None:
            outline_level = OxmlElement("w:outlineLvl")
            paragraph_properties.append(outline_level)
        outline_level.set(qn("w:val"), str(level - 1))

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_complex_field(paragraph, "PAGE", "1")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(
    paragraph,
    text: str,
    anchor: str,
    style: "InlineStyle",
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = paragraph.add_run(text)
    run.bold = style.bold
    run.italic = style.italic
    run.font.color.rgb = RGBColor(17, 85, 204)
    run.font.underline = True
    if style.code:
        set_run_font(run, CODE_FONT, 9)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def normalized_prose(text: str | None) -> str:
    return re.sub(r"\s+", " ", text or "")


def available_paragraph_style(document: Document, preferred: str) -> str:
    names = {style.name for style in document.styles}
    if preferred in names:
        return preferred
    for fallback in ("Body Text", "normal", "Normal"):
        if fallback in names:
            return fallback
    raise ValueError("Base DOCX has no usable paragraph style")


@dataclass(frozen=True)
class InlineStyle:
    bold: bool = False
    italic: bool = False
    code: bool = False


def _add_plain_text(paragraph, value: str, style: InlineStyle) -> None:
    if not value:
        return
    run = paragraph.add_run(value)
    run.bold = style.bold
    run.italic = style.italic
    if style.code:
        set_run_font(run, CODE_FONT, 9)
        run.font.color.rgb = RGBColor(55, 55, 55)


def _reference_anchor(label: str, number: str) -> str:
    normalized = label.casefold()
    if normalized.startswith("глав"):
        prefix = "ch"
    elif normalized.startswith("рисун"):
        prefix = "fig"
    elif normalized.startswith("таблиц"):
        prefix = "table"
    else:
        prefix = "listing"
    return f"{prefix}_{number}"


def add_text(
    paragraph,
    value: str,
    style: InlineStyle,
    internal_anchors: set[str] | None = None,
    link_internal_references: bool = True,
) -> None:
    if not value:
        return
    if not paragraph.text and not paragraph._p.findall(qn("w:hyperlink")):
        value = value.lstrip()
    if not link_internal_references or style.code or not internal_anchors:
        _add_plain_text(paragraph, value, style)
        return

    position = 0
    for match in INTERNAL_REFERENCE_RE.finditer(value):
        anchor = _reference_anchor(match.group("label"), match.group("number"))
        if anchor not in internal_anchors:
            continue
        _add_plain_text(paragraph, value[position : match.start()], style)
        add_internal_hyperlink(paragraph, match.group(0), anchor, style)
        position = match.end()
    _add_plain_text(paragraph, value[position:], style)


def render_inline(
    paragraph,
    element,
    style: InlineStyle = InlineStyle(),
    internal_anchors: set[str] | None = None,
    link_internal_references: bool = True,
) -> None:
    add_text(
        paragraph,
        normalized_prose(element.text),
        style,
        internal_anchors,
        link_internal_references,
    )
    for child in element:
        tag = child.tag.lower() if isinstance(child.tag, str) else ""
        if tag in {"ul", "ol", "table", "pre"}:
            continue
        if tag == "strong":
            render_inline(
                paragraph,
                child,
                InlineStyle(True, style.italic, style.code),
                internal_anchors,
                link_internal_references,
            )
        elif tag == "em":
            render_inline(
                paragraph,
                child,
                InlineStyle(style.bold, True, style.code),
                internal_anchors,
                link_internal_references,
            )
        elif tag == "code":
            render_inline(
                paragraph,
                child,
                InlineStyle(style.bold, style.italic, True),
                internal_anchors,
                False,
            )
        elif tag == "a":
            add_hyperlink(paragraph, "".join(child.itertext()).strip(), child.get("href", ""))
        elif tag == "br":
            paragraph.add_run().add_break()
        elif tag != "img":
            render_inline(
                paragraph,
                child,
                style,
                internal_anchors,
                link_internal_references,
            )
        add_text(
            paragraph,
            normalized_prose(child.tail),
            style,
            internal_anchors,
            link_internal_references,
        )


class BookmarkManager:
    def __init__(self, document: Document) -> None:
        identifiers = [
            int(node.get(qn("w:id")))
            for node in document._element.findall(f".//{qn('w:bookmarkStart')}")
        ]
        self.next_id = max(identifiers, default=0) + 1
        self.names = {
            node.get(qn("w:name"))
            for node in document._element.findall(f".//{qn('w:bookmarkStart')}")
        }

    def add(self, paragraph, name: str) -> None:
        if name in self.names:
            raise ValueError(f"Duplicate bookmark name: {name}")
        identifier = str(self.next_id)
        self.next_id += 1
        self.names.add(name)

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), identifier)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), identifier)

        insertion_index = 1 if paragraph._p.pPr is not None else 0
        paragraph._p.insert(insertion_index, start)
        paragraph._p.append(end)


class NumberingManager:
    def __init__(self, document: Document) -> None:
        self.root = document.part.numbering_part.element
        self.abstract_ids = {kind: self._find_abstract_id(kind) for kind in ("bullet", "decimal")}
        self.next_num_id = (
            max([int(node.get(qn("w:numId"))) for node in self.root.findall(qn("w:num"))] or [0])
            + 1
        )

    def _find_abstract_id(self, kind: str) -> int:
        for abstract in self.root.findall(qn("w:abstractNum")):
            level = abstract.find(qn("w:lvl"))
            if level is None:
                continue
            number_format = level.find(qn("w:numFmt"))
            if number_format is not None and number_format.get(qn("w:val")) == kind:
                return int(abstract.get(qn("w:abstractNumId")))
        raise ValueError(f"Base DOCX has no {kind} numbering definition")

    def new_list(self, kind: str) -> int:
        number_id = self.next_num_id
        self.next_num_id += 1
        number = OxmlElement("w:num")
        number.set(qn("w:numId"), str(number_id))
        abstract = OxmlElement("w:abstractNumId")
        abstract.set(qn("w:val"), str(self.abstract_ids[kind]))
        number.append(abstract)
        self.root.append(number)
        return number_id

    @staticmethod
    def apply(paragraph, number_id: int, level: int) -> None:
        properties = paragraph._p.get_or_add_pPr()
        number_properties = properties.find(qn("w:numPr"))
        if number_properties is None:
            number_properties = OxmlElement("w:numPr")
            properties.append(number_properties)
        indentation = OxmlElement("w:ilvl")
        indentation.set(qn("w:val"), str(min(level, 8)))
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(number_id))
        number_properties.extend((indentation, number))


class DocxRenderer:
    def __init__(self, document: Document, manuscript: Path) -> None:
        self.document = document
        self.manuscript = manuscript
        self.numbering = NumberingManager(document)
        self.bookmarks = BookmarkManager(document)
        self.available_anchors: set[str] = set()
        section = document.sections[0]
        self.usable_width = section.page_width - section.left_margin - section.right_margin
        self.usable_height = section.page_height - section.top_margin - section.bottom_margin
        self.metrics = {
            "headings": 0,
            "paragraphs": 0,
            "lists": 0,
            "list_items": 0,
            "tables": 0,
            "images": 0,
            "code_lines": 0,
            "page_breaks": 0,
            "bookmarks": 0,
            "internal_hyperlinks": 0,
        }
        self.is_first_element = True
        self.in_chapter_sources = False
        self.toc_entries: list[tuple[int, str]] = []

    def render(self, root) -> None:
        self.available_anchors = {
            anchor
            for element in root
            if (anchor := self._bookmark_anchor(element)) is not None
        }
        self.toc_entries = [
            (int(element.tag[1]), normalized_prose("".join(element.itertext())).strip())
            for element in root
            if isinstance(element.tag, str)
            and element.tag.lower() in {"h1", "h2"}
            and self._include_in_static_toc(
                int(element.tag[1]),
                normalized_prose("".join(element.itertext())).strip(),
            )
        ]
        for element in root:
            self.render_block(element)

    @staticmethod
    def _bookmark_anchor(element) -> str | None:
        if not isinstance(element.tag, str):
            return None
        tag = element.tag.lower()
        if tag not in {"h2", "h3", "p"}:
            return None
        value = normalized_prose("".join(element.itertext())).strip()
        patterns = (
            (r"Глава (\d+)\. .+", "ch"),
            (r"Лабораторная работа (\d+)\. .+", "lab"),
            (r"Рисунок (\d+)\. .+", "fig"),
            (r"Таблица (\d+)\. .+", "table"),
            (r"Листинг (\d+)\. .+", "listing"),
        )
        for pattern, prefix in patterns:
            if match := re.fullmatch(pattern, value):
                return f"{prefix}_{match.group(1)}"
        return None

    def _render_inline(self, paragraph, element, *, allow_internal_links: bool = True) -> None:
        before = len(paragraph._p.findall(qn("w:hyperlink")))
        render_inline(
            paragraph,
            element,
            internal_anchors=self.available_anchors,
            link_internal_references=allow_internal_links,
        )
        after = len(paragraph._p.findall(qn("w:hyperlink")))
        if allow_internal_links:
            self.metrics["internal_hyperlinks"] += sum(
                1
                for node in paragraph._p.findall(qn("w:hyperlink"))[before:after]
                if node.get(qn("w:anchor"))
            )

    def _add_bookmark(self, paragraph, element) -> None:
        anchor = self._bookmark_anchor(element)
        if anchor is None:
            return
        self.bookmarks.add(paragraph, anchor)
        self.metrics["bookmarks"] += 1

    @staticmethod
    def _include_in_static_toc(level: int, title: str) -> bool:
        if level == 1:
            return True
        return title.startswith(
            (
                "Об авторе",
                "Как использовать примеры безопасно",
                "Введение.",
                "Глава ",
                "Итоговый проект.",
                "Как пользоваться приложениями",
                "Приложение ",
            )
        )

    def add_paragraph(self):
        paragraph = self.document.add_paragraph()
        self.metrics["paragraphs"] += 1
        return paragraph

    def render_block(self, element, list_level: int = 0) -> None:
        tag = element.tag.lower() if isinstance(element.tag, str) else ""
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self.render_heading(element, int(tag[1]))
        elif tag == "p":
            self._render_paragraph(element)
        elif tag in {"ul", "ol"}:
            self.render_list(element, list_level)
        elif tag == "pre":
            self.render_code_block(element)
        elif tag == "table":
            self.render_table(element)
        elif tag == "blockquote":
            self._render_blockquote(element)
        elif tag == "hr":
            self._render_horizontal_rule()
        else:
            for child in element:
                self.render_block(child, list_level)

    def _render_paragraph(self, element) -> None:
        images = element.findall("img")
        prose = normalized_prose(element.text).strip()
        if len(images) == 1 and not prose and len(element) == 1:
            self.render_image(images[0])
        else:
            paragraph = self.add_paragraph()
            if self.is_first_element:
                self.render_title(paragraph, element)
            else:
                self._render_inline(
                    paragraph,
                    element,
                    allow_internal_links=self._bookmark_anchor(element) is None,
                )
                self._add_bookmark(paragraph, element)
                if re.fullmatch(r"Таблица \d+\. .+", paragraph.text.strip()):
                    paragraph.style = available_paragraph_style(
                        self.document,
                        "Caption",
                    )
                    set_keep_next(paragraph)
                    set_keep_lines(paragraph)
                if re.match(r"^S\d{3}\.", paragraph.text.strip()):
                    paragraph.paragraph_format.left_indent = Inches(0.28)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.28)
                    paragraph.paragraph_format.space_after = Pt(2)
                    set_keep_lines(paragraph)
                elif self.in_chapter_sources:
                    set_keep_next(paragraph)
            self.is_first_element = False

    def _render_blockquote(self, element) -> None:
        for child in element:
            paragraph = self.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            render_inline(
                paragraph,
                child,
                InlineStyle(italic=True),
                self.available_anchors,
            )

    def _render_horizontal_rule(self) -> None:
        paragraph = self.add_paragraph()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "B7B7B7")
        border.append(bottom)
        paragraph._p.get_or_add_pPr().append(border)

    def render_title(self, paragraph, element) -> None:
        render_inline(
            paragraph,
            element,
            internal_anchors=self.available_anchors,
            link_internal_references=False,
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        set_keep_next(paragraph)
        for run in paragraph.runs:
            set_run_font(run, "Arial", 26)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0)
        self.render_table_of_contents()

    def render_table_of_contents(self) -> None:
        if not self.toc_entries:
            raise ValueError("Manuscript has no headings for the table of contents")
        body_style = available_paragraph_style(self.document, "Body Text")
        heading = self.document.add_paragraph(style=body_style)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)
        set_keep_next(heading)
        run = heading.add_run("Оглавление")
        set_run_font(run, "Arial", 16)
        run.bold = True

        result_paragraphs = []
        for index, (level, title) in enumerate(self.toc_entries):
            paragraph = self.document.add_paragraph(style=body_style)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            if level == 2:
                paragraph.paragraph_format.left_indent = Inches(0.24)
            if index == 0:
                begin_complex_field(paragraph, 'TOC \\o "1-2" \\h \\z \\u')
            run = paragraph.add_run(title)
            run.bold = level == 1
            if level == 1:
                set_keep_next(paragraph)
            result_paragraphs.append(paragraph)
        end_complex_field(result_paragraphs[-1])

        page_break = self.document.add_paragraph(style=body_style)
        page_break.add_run().add_break(WD_BREAK.PAGE)
        self.metrics["paragraphs"] += len(result_paragraphs) + 2

    def render_heading(self, element, level: int) -> None:
        paragraph = self.document.add_paragraph(style=f"Heading {min(level, 6)}")
        render_inline(
            paragraph,
            element,
            internal_anchors=self.available_anchors,
            link_internal_references=False,
        )
        self._add_bookmark(paragraph, element)
        self.metrics["headings"] += 1
        text = paragraph.text.strip()
        if self.in_chapter_sources and text != "Источники главы":
            self.in_chapter_sources = False
        if text == "Источники главы":
            self.in_chapter_sources = True
        page_break = level == 1
        if level == 2:
            match = re.fullmatch(r"Глава (\d+)\. .+", text)
            page_break = match is not None and int(match.group(1)) not in FIRST_CHAPTERS_IN_PART
        if page_break:
            set_page_break_before(paragraph)
            self.metrics["page_breaks"] += 1
        set_keep_next(paragraph)
        self.is_first_element = False

    def render_list(self, element, level: int) -> None:
        kind = "bullet" if element.tag.lower() == "ul" else "decimal"
        number_id = self.numbering.new_list(kind)
        self.metrics["lists"] += 1
        items = element.findall("li")
        for index, item in enumerate(items):
            paragraph = self.add_paragraph()
            self.numbering.apply(paragraph, number_id, level)
            self._render_inline(paragraph, item)
            if self.in_chapter_sources and index < len(items) - 1:
                set_keep_next(paragraph)
            self.metrics["list_items"] += 1
            for nested in item:
                if isinstance(nested.tag, str) and nested.tag.lower() in {"ul", "ol"}:
                    self.render_list(nested, level + 1)
        if self.in_chapter_sources:
            self.in_chapter_sources = False

    def render_code_block(self, element) -> None:
        code = element.find("code")
        value = "".join(code.itertext()) if code is not None else "".join(element.itertext())
        for line in value.rstrip("\n").split("\n"):
            paragraph = self.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            set_keep_lines(paragraph)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F5F5F5")
            paragraph._p.get_or_add_pPr().append(shading)
            run = paragraph.add_run(line or " ")
            set_run_font(run, CODE_FONT, 8.5)
            self.metrics["code_lines"] += 1

    def render_image(self, element) -> None:
        source = element.get("src", "")
        image_path = (self.manuscript.parent / source).resolve()
        if not image_path.is_file():
            raise FileNotFoundError(f"Missing manuscript image: {source}")
        with Image.open(image_path) as image:
            width_px, height_px = image.size
            dpi = image.info.get("dpi", (96, 96))
            dpi_x = float(dpi[0] or 96)
            dpi_y = float(dpi[1] or 96)
        natural_width = Inches(width_px / dpi_x)
        natural_height = Inches(height_px / dpi_y)
        scale = min(
            1.0,
            self.usable_width / natural_width,
            (self.usable_height * 0.82) / natural_height,
        )
        paragraph = self.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_keep_next(paragraph)
        shape = paragraph.add_run().add_picture(str(image_path), width=int(natural_width * scale))
        alt = element.get("alt", "").strip() or image_path.stem
        shape._inline.docPr.set("title", "Иллюстрация к рукописи")
        shape._inline.docPr.set("descr", alt[:250])
        self.metrics["images"] += 1

    def render_table(self, element) -> None:
        rows = element.xpath("./thead/tr | ./tbody/tr | ./tr")
        if not rows:
            return
        column_count = max(len(row.xpath("./th | ./td")) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=column_count)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)
        text_lengths = [1] * column_count
        for row_index, row in enumerate(rows):
            row_properties = table.rows[row_index]._tr.get_or_add_trPr()
            keep_row_together = OxmlElement("w:cantSplit")
            keep_row_together.set(qn("w:val"), "true")
            row_properties.append(keep_row_together)
            cells = row.xpath("./th | ./td")
            for column_index, cell in enumerate(cells):
                text_length = self._render_table_cell(
                    table.cell(row_index, column_index),
                    cell,
                    row_index,
                    len(rows),
                )
                text_lengths[column_index] = max(
                    text_lengths[column_index],
                    text_length,
                )
        shares = self._table_column_shares(rows, text_lengths)
        self._apply_table_column_widths(table, shares)
        self.metrics["tables"] += 1

    def _render_table_cell(self, target, cell, row_index: int, row_count: int) -> int:
        target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = target.paragraphs[0]
        render_inline(
            paragraph,
            cell,
            InlineStyle(bold=row_index == 0 or cell.tag.lower() == "th"),
            self.available_anchors,
        )
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if row_index in {0, 1, max(0, row_count - 2)}:
            set_keep_next(paragraph)
        for run in paragraph.runs:
            fonts = run._r.get_or_add_rPr().rFonts
            if fonts is not None and any(
                "mono" in value.lower() or "courier" in value.lower()
                for value in fonts.attrib.values()
            ):
                run.font.size = Pt(8)
        if row_index == 0:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EDEDED")
            target._tc.get_or_add_tcPr().append(shading)
        return len(paragraph.text)

    @staticmethod
    def _table_column_shares(rows, text_lengths: list[int]) -> list[float]:
        column_count = len(text_lengths)
        weights = [min(length, 48) for length in text_lengths]
        total_weight = sum(weights)
        minimum_share = min(0.15, 0.8 / column_count)
        proportional_share = 1 - minimum_share * column_count
        shares = [
            minimum_share + proportional_share * weight / total_weight
            for weight in weights
        ]
        first_column_values = [
            "".join(rows[row_index].xpath("./th | ./td")[0].itertext())
            for row_index in range(len(rows))
            if rows[row_index].xpath("./th | ./td")
        ]
        if column_count >= 3 and any(
            re.search(r"\b[a-z][a-z0-9]*(?:_[a-z0-9]+){2,}\b", value)
            for value in first_column_values
        ):
            requested = max(shares[0], 0.32)
            remainder = 1 - requested
            other_total = sum(shares[1:])
            shares = [requested, *[share * remainder / other_total for share in shares[1:]]]
        return shares

    def _apply_table_column_widths(self, table, shares: list[float]) -> None:
        for column_index, share in enumerate(shares):
            width = int(self.usable_width * share)
            table.columns[column_index].width = width
            for cell in table.columns[column_index].cells:
                cell.width = width


def prune_unused_document_relationships(document: Document) -> None:
    used = {
        value
        for node in document._element.iter()
        for attribute, value in node.attrib.items()
        if attribute in {qn("r:id"), qn("r:embed"), qn("r:link")}
    }
    for relationship_id, relationship in list(document.part.rels.items()):
        if (
            relationship.reltype
            in {
                RELATIONSHIP_TYPE.IMAGE,
                RELATIONSHIP_TYPE.HYPERLINK,
            }
            and relationship_id not in used
        ):
            document.part.drop_rel(relationship_id)


def build(
    manuscript: Path,
    base_docx: Path,
    output_docx: Path,
    node: Path,
    node_path: Path,
) -> dict[str, object]:
    source_html = markdown_to_html(manuscript, node, node_path)
    root = html.fragment_fromstring(source_html, create_parent="div")
    document = Document(str(base_docx))
    clear_document_body(document)
    configure_document_properties(document)
    renderer = DocxRenderer(document, manuscript)
    renderer.render(root)
    prune_unused_document_relationships(document)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return {
        "manuscript": str(manuscript),
        "base_docx": str(base_docx),
        "output_docx": str(output_docx),
        "output_bytes": output_docx.stat().st_size,
        "output_sha256": sha256(output_docx),
        **renderer.metrics,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Russian editorial DOCX")
    parser.add_argument("--manuscript", required=True, type=Path)
    parser.add_argument("--base-docx", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    parser.add_argument("--node", required=True, type=Path)
    parser.add_argument("--node-path", required=True, type=Path)
    parser.add_argument("--metrics-json", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    metrics = build(
        args.manuscript,
        args.base_docx,
        args.output_docx,
        args.node,
        args.node_path,
    )
    payload = json.dumps(metrics, ensure_ascii=False, indent=2, sort_keys=True)
    if args.metrics_json:
        args.metrics_json.write_text(payload + "\n", encoding="utf-8")
    print(payload)


if __name__ == "__main__":
    main()
