#!/usr/bin/env python3
"""Build a macro-free Template2000n DOCX derivative from a raw Google Docs DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import zipfile
from collections import Counter
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from xml.etree import ElementTree as ET

from docx import Document
from PIL import Image

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}
for prefix, uri in NS.items():
    ET.register_namespace(prefix, uri)


STYLE_PARTS = {
    "word/styles.xml",
    "word/theme/theme1.xml",
}

EMBEDDED_FONT_TAGS = {
    f"{{{NS['w']}}}embedRegular",
    f"{{{NS['w']}}}embedBold",
    f"{{{NS['w']}}}embedItalic",
    f"{{{NS['w']}}}embedBoldItalic",
}

FONT_ATTRIBUTES = {
    f"{{{NS['w']}}}ascii",
    f"{{{NS['w']}}}hAnsi",
    f"{{{NS['w']}}}eastAsia",
    f"{{{NS['w']}}}cs",
}

CALLOUT_HEADING_LABELS = frozenset(
    {
        "Практическая проверка.",
        "Практическая проверка в репозитории.",
        "Связь со следующей главой.",
        "Сопутствующие материалы.",
        "Частые ошибки.",
        "Граница доказательств.",
    }
)

PRESERVED_PARAGRAPH_STYLE_IDS = frozenset(
    {
        "Heading1",
        "Heading2",
        "Heading3",
        "Heading4",
        "Heading5",
        "Title",
    }
)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def paragraph_texts(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(str(path)).paragraphs]


def document_text_nodes(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    return [node.text or "" for node in root.iter(f"{{{NS['w']}}}t")]


def media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }


def flatten_alpha_images(media_dir: Path) -> set[str]:
    flattened: set[str] = set()
    for path in sorted(media_dir.glob("*.png")):
        source = path.read_bytes()
        with Image.open(BytesIO(source)) as image:
            if "A" not in image.mode and "transparency" not in image.info:
                continue
            rgba = image.convert("RGBA")
            output = Image.new("RGB", rgba.size, "white")
            output.paste(rgba, mask=rgba.getchannel("A"))
            save_options: dict[str, object] = {"optimize": True}
            if dpi := image.info.get("dpi"):
                save_options["dpi"] = dpi
            if icc_profile := image.info.get("icc_profile"):
                save_options["icc_profile"] = icc_profile
            output.save(path, "PNG", **save_options)
        flattened.add(f"word/media/{path.name}")
    return flattened


def word_count(texts: list[str]) -> int:
    return sum(len(text.split()) for text in texts)


def style_counts(path: Path) -> Counter[str]:
    counts: Counter[str] = Counter()
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        style = paragraph.style.name if paragraph.style is not None else "<none>"
        counts[style] += 1
    return counts


def remove_heading_numbering(styles_xml: bytes) -> bytes:
    root = ET.fromstring(styles_xml)
    for style in root.findall("w:style", NS):
        style_id = style.get(f"{{{NS['w']}}}styleId", "")
        if style_id not in {"Heading1", "Heading2", "Heading3", "Heading4", "Heading5"}:
            continue
        ppr = style.find("w:pPr", NS)
        if ppr is None:
            continue
        for numpr in list(ppr.findall("w:numPr", NS)):
            ppr.remove(numpr)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def normalize_numbering_fonts(numbering_xml: bytes) -> tuple[bytes, int]:
    root = ET.fromstring(numbering_xml)
    replacements = 0
    for fonts in root.findall(".//w:rFonts", NS):
        for attribute in FONT_ATTRIBUTES:
            if fonts.get(attribute) == "Noto Sans Symbols":
                fonts.set(attribute, "Times New Roman")
                replacements += 1
    return ET.tostring(root, encoding="utf-8", xml_declaration=True), replacements


def registered_font_relationships(font_table_xml: bytes) -> set[str]:
    root = ET.fromstring(font_table_xml)
    relationship_attr = f"{{{NS['r']}}}id"
    return {
        relationship_id
        for node in root.iter()
        if node.tag in EMBEDDED_FONT_TAGS
        if (relationship_id := node.get(relationship_attr))
    }


def merge_preserved_styles(template_xml: bytes, raw_xml: bytes) -> tuple[bytes, int]:
    template = ET.fromstring(template_xml)
    raw = ET.fromstring(raw_xml)
    style_id_attribute = f"{{{NS['w']}}}styleId"
    existing = {style.get(style_id_attribute) for style in template.findall("w:style", NS)}
    preserved = 0
    for style in raw.findall("w:style", NS):
        style_id = style.get(style_id_attribute)
        if style_id != "Title" or style_id in existing:
            continue
        template.append(ET.fromstring(ET.tostring(style)))
        preserved += 1
    return ET.tostring(template, encoding="utf-8", xml_declaration=True), preserved


def paragraph_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()


def set_paragraph_style(paragraph: ET.Element, style_id: str) -> None:
    ppr = paragraph.find("w:pPr", NS)
    if ppr is None:
        ppr = ET.Element(f"{{{NS['w']}}}pPr")
        paragraph.insert(0, ppr)
    pstyle = ppr.find("w:pStyle", NS)
    if pstyle is None:
        pstyle = ET.Element(f"{{{NS['w']}}}pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(f"{{{NS['w']}}}val", style_id)


def set_paragraph_flag(paragraph: ET.Element, name: str) -> None:
    ppr = paragraph.find("w:pPr", NS)
    if ppr is None:
        ppr = ET.Element(f"{{{NS['w']}}}pPr")
        paragraph.insert(0, ppr)
    if ppr.find(f"w:{name}", NS) is None:
        ppr.append(ET.Element(f"{{{NS['w']}}}{name}"))


def paragraph_is_monospace(paragraph: ET.Element) -> bool:
    styled_characters = 0
    monospace_characters = 0
    for run in paragraph.findall("w:r", NS):
        run_text = "".join(node.text or "" for node in run.findall(".//w:t", NS))
        if not run_text.strip():
            continue
        fonts = run.find("w:rPr/w:rFonts", NS)
        if fonts is None:
            continue
        styled_characters += len(run_text)
        if any(
            "mono" in value.lower() or "courier" in value.lower() for value in fonts.attrib.values()
        ):
            monospace_characters += len(run_text)
    return styled_characters > 0 and monospace_characters / styled_characters >= 0.8


def _table_paragraph_styles(root: ET.Element) -> dict[ET.Element, str]:
    styles: dict[ET.Element, str] = {}
    for table in root.findall(".//w:tbl", NS):
        for row_index, row in enumerate(table.findall("w:tr", NS)):
            style_id = "Style20" if row_index == 0 else "Style21"
            for paragraph in row.findall(".//w:p", NS):
                styles[paragraph] = style_id
    return styles


def _next_nonempty_paragraph_text(
    paragraphs: list[ET.Element],
    index: int,
) -> str:
    return next(
        (
            candidate_text
            for candidate in paragraphs[index + 1 :]
            if (candidate_text := paragraph_text(candidate))
        ),
        "",
    )


def _style_image_paragraph(
    paragraph: ET.Element,
    paragraphs: list[ET.Element],
    index: int,
    last_text: str,
    mappings: Counter[str],
) -> None:
    set_paragraph_style(paragraph, "Style28")
    mappings["picture"] += 1
    following_text = _next_nonempty_paragraph_text(paragraphs, index)
    if re.match(r"^Рисунок \d+\.", following_text):
        set_paragraph_flag(paragraph, "keepNext")
        mappings["picture_kept_with_caption"] += 1
    fallback_description = (following_text or last_text)[:250]
    for properties in paragraph.findall(".//wp:docPr", NS):
        properties.set("title", "Иллюстрация к рукописи")
        if not properties.get("descr", "").strip():
            properties.set("descr", fallback_description)
        mappings["image_alt_text"] += 1


def _semantic_paragraph_style(
    paragraph: ET.Element,
    text: str,
    table_styles: dict[ET.Element, str],
    pending_callout_body: bool,
    mappings: Counter[str],
) -> tuple[str, bool]:
    if paragraph in table_styles:
        target_style = table_styles[paragraph]
        mappings["table_header" if target_style == "Style20" else "table_body"] += 1
    elif text in CALLOUT_HEADING_LABELS:
        target_style = "Style24"
        mappings["callout_heading"] += 1
        pending_callout_body = True
    elif pending_callout_body:
        target_style = "Style23"
        mappings["callout_body"] += 1
        pending_callout_body = False
    elif re.match(r"^Таблица \d+\.", text):
        target_style = "Style17"
        set_paragraph_flag(paragraph, "keepNext")
        set_paragraph_flag(paragraph, "keepLines")
        mappings["table_caption"] += 1
    elif re.match(r"^Рисунок \d+\.", text):
        target_style = "Style17"
        set_paragraph_flag(paragraph, "keepLines")
        mappings["figure_caption"] += 1
    elif paragraph.find("w:pPr/w:numPr", NS) is not None:
        target_style = "Style18"
        mappings["list"] += 1
    elif paragraph_is_monospace(paragraph):
        target_style = "Style16"
        mappings["program"] += 1
    else:
        target_style = "BodyText"
        mappings["body_text"] += 1
    return target_style, pending_callout_body


def map_semantic_styles(document_xml: bytes) -> tuple[bytes, Counter[str]]:
    root = ET.fromstring(document_xml)
    paragraphs = root.findall(".//w:p", NS)
    mappings: Counter[str] = Counter()
    table_styles = _table_paragraph_styles(root)
    pending_callout_body = False
    last_text = "Схема архитектуры безопасного ИИ-агента"
    style_attribute = f"{{{NS['w']}}}val"

    for index, paragraph in enumerate(paragraphs):
        text = paragraph_text(paragraph)
        current_style = paragraph.find("w:pPr/w:pStyle", NS)
        current_style_id = current_style.get(style_attribute) if current_style is not None else None
        has_image = (
            paragraph.find(".//w:drawing", NS) is not None
            or paragraph.find(".//w:pict", NS) is not None
        )

        if has_image:
            _style_image_paragraph(paragraph, paragraphs, index, last_text, mappings)
            pending_callout_body = False
            continue

        if not text:
            continue
        if current_style_id in PRESERVED_PARAGRAPH_STYLE_IDS:
            pending_callout_body = False
            last_text = text
            continue

        target_style, pending_callout_body = _semantic_paragraph_style(
            paragraph,
            text,
            table_styles,
            pending_callout_body,
            mappings,
        )

        set_paragraph_style(paragraph, target_style)
        last_text = text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True), mappings


def build(raw_docx: Path, template_docx: Path, output_docx: Path) -> dict[str, object]:
    with TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        raw_dir = tmpdir / "raw"
        template_dir = tmpdir / "template"
        raw_dir.mkdir()
        template_dir.mkdir()

        with zipfile.ZipFile(raw_docx) as zf:
            zf.extractall(raw_dir)
        with zipfile.ZipFile(template_docx) as zf:
            zf.extractall(template_dir)

        raw_styles = (raw_dir / "word/styles.xml").read_bytes()
        title_styles_preserved = 0
        for part in STYLE_PARTS:
            src = template_dir / part
            dst = raw_dir / part
            if not src.exists():
                raise FileNotFoundError(f"Template part missing: {part}")
            if part == "word/styles.xml":
                merged_styles, title_styles_preserved = merge_preserved_styles(
                    src.read_bytes(), raw_styles
                )
                dst.write_bytes(remove_heading_numbering(merged_styles))
            else:
                dst.parent.mkdir(parents=True, exist_ok=True)
                shutil.copyfile(src, dst)

        numbering = raw_dir / "word" / "numbering.xml"
        normalized_numbering, numbering_font_replacements = normalize_numbering_fonts(
            numbering.read_bytes()
        )
        numbering.write_bytes(normalized_numbering)

        font_table = raw_dir / "word" / "fontTable.xml"
        embedded_font_registrations = len(registered_font_relationships(font_table.read_bytes()))

        document_xml = raw_dir / "word" / "document.xml"
        mapped_xml, semantic_mappings = map_semantic_styles(document_xml.read_bytes())
        document_xml.write_bytes(mapped_xml)

        flattened_media = flatten_alpha_images(raw_dir / "word" / "media")

        if output_docx.exists():
            output_docx.unlink()
        with zipfile.ZipFile(output_docx, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in sorted(raw_dir.rglob("*")):
                if path.is_file():
                    zf.write(path, path.relative_to(raw_dir).as_posix())

    raw_texts = paragraph_texts(raw_docx)
    output_texts = paragraph_texts(output_docx)
    equal_text = raw_texts == output_texts
    if not equal_text:
        raise AssertionError("Raw and Template2000n paragraph text differ")

    raw_text_nodes = document_text_nodes(raw_docx)
    output_text_nodes = document_text_nodes(output_docx)
    equal_text_nodes = raw_text_nodes == output_text_nodes
    if not equal_text_nodes:
        raise AssertionError("Raw and Template2000n document text nodes differ")

    raw_media = media_hashes(raw_docx)
    output_media = media_hashes(output_docx)
    equal_media_files = raw_media.keys() == output_media.keys()
    equal_untouched_media = all(
        raw_hash == output_media[name]
        for name, raw_hash in raw_media.items()
        if name not in flattened_media
    )
    if not equal_media_files or not equal_untouched_media:
        raise AssertionError("Template2000n media changed outside alpha normalization")

    raw_counts = style_counts(raw_docx)
    output_counts = style_counts(output_docx)
    return {
        "raw_docx": str(raw_docx),
        "template_docx": str(template_docx),
        "output_docx": str(output_docx),
        "output_bytes": output_docx.stat().st_size,
        "output_sha256": sha256(output_docx),
        "paragraphs": len(output_texts),
        "non_empty_paragraphs": sum(1 for text in output_texts if text.strip()),
        "document_text_nodes": len(output_text_nodes),
        "approximate_words": word_count(output_texts),
        "bodytext_mapped": semantic_mappings["body_text"],
        "semantic_style_mappings": dict(semantic_mappings),
        "title_styles_preserved": title_styles_preserved,
        "embedded_font_registrations_preserved": embedded_font_registrations,
        "numbering_font_attributes_normalized": numbering_font_replacements,
        "text_equality": equal_text,
        "document_text_equality": equal_text_nodes,
        "media_files": len(output_media),
        "media_equality": equal_media_files and equal_untouched_media,
        "media_byte_equality": raw_media == output_media,
        "alpha_images_flattened": len(flattened_media),
        "raw_style_counts": dict(raw_counts),
        "output_style_counts": dict(output_counts),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--raw-docx", required=True, type=Path)
    parser.add_argument("--template-docx", required=True, type=Path)
    parser.add_argument("--output-docx", required=True, type=Path)
    parser.add_argument("--metrics-json", type=Path)
    args = parser.parse_args()

    metrics = build(args.raw_docx, args.template_docx, args.output_docx)
    payload = json.dumps(metrics, ensure_ascii=False, indent=2, sort_keys=True)
    if args.metrics_json:
        args.metrics_json.write_text(payload + "\n", encoding="utf-8")
    print(payload)


if __name__ == "__main__":
    main()
