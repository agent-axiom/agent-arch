import ast
import copy
import hashlib
import json
import os
import posixpath
import re
import subprocess
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import pytest

from docs.publisher.tools import (
    build_ru_editorial_docx,
    generate_publisher_layout_v2,
    normalize_docx_figure_caption_order,
    render_qa_metrics,
    sync_ru_docx_visuals,
)

ROOT = Path(__file__).resolve().parents[1]
EDITORIAL_BUILDER = ROOT / "docs/publisher/tools/build_ru_editorial_docx.py"
RAW_EDITORIAL_DOCX = ROOT / (
    "docs/publisher/artifacts/agent-arch-ru-google-doc-publication-readiness-2026-08-14.docx"
)
EDITORIAL_TEMPLATE_DOCX = ROOT / (
    "docs/publisher/artifacts/agent-arch-ru-template2000n-publication-readiness-2026-08-14.docx"
)
EDITORIAL_MANUSCRIPT = ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md"
CURRENT_RAW_DOCX = ROOT / (
    "docs/publisher/artifacts/agent-arch-ru-google-doc-book-standards-2026-08-23.docx"
)
CURRENT_TEMPLATE_DOCX = ROOT / (
    "docs/publisher/artifacts/agent-arch-ru-template2000n-book-standards-2026-08-23.docx"
)
RAW_A11Y = ROOT / "docs/publisher/ru-google-doc-book-standards-2026-08-23.a11y.json"
RAW_RENDER_QA = ROOT / "docs/publisher/ru-google-doc-book-standards-2026-08-23.render-qa.json"
TEMPLATE_A11Y = ROOT / "docs/publisher/ru-template2000n-book-standards-2026-08-23.a11y.json"
TEMPLATE_FONT_AUDIT = ROOT / (
    "docs/publisher/ru-template2000n-book-standards-2026-08-23.font-audit.json"
)
TEMPLATE_METRICS = ROOT / (
    "docs/publisher/ru-template2000n-book-standards-2026-08-23.metrics.json"
)
TEMPLATE_RENDER_QA = ROOT / (
    "docs/publisher/ru-template2000n-book-standards-2026-08-23.render-qa.json"
)
TEMPLATE_VISUAL_AUDIT = ROOT / (
    "docs/publisher/ru-template2000n-book-standards-2026-08-23.visual-audit.json"
)
VISUAL_LAYOUT_AUDIT = ROOT / "docs/publisher/ru-visual-layout-audit-2026-08-24.json"
INLINE_DIAGRAMS = ROOT / "docs/publisher/ru-inline-diagrams-2026-07-13.json"
NUMBERED_DIAGRAMS = ROOT / "docs/publisher/ru-numbered-diagrams-2026-07-15.json"
EDITORIAL_DIAGRAMS = ROOT / "docs/publisher/ru-editorial-diagrams-2026-07-16.json"
LAYOUT_V2_INVENTORY = ROOT / "docs/publisher/ru-publisher-layout-v2-inventory.json"
LAYOUT_V2_LEDGER = ROOT / "docs/publisher/ru-publisher-layout-v2-review-ledger.json"
FINAL_LAYOUT_DOCX = ROOT / (
    "docs/publisher/artifacts/agent-arch-ru-template2000n-publisher-layout-v2-2026-08-30.docx"
)
FINAL_LAYOUT_RENDER_QA = ROOT / "docs/publisher/qa/layout-v2/final-docx/render-qa.json"
FINAL_LAYOUT_VISUAL_AUDIT = ROOT / "docs/publisher/qa/layout-v2/final-docx/visual-audit.json"
TASK_3A_QA_DIR = ROOT / "docs/publisher/qa/layout-v2/task-3a"
TASK_3A_RENDERER_REPORT = TASK_3A_QA_DIR / "renderer-report.json"
TASK_3A_CONTACT_SHEET = TASK_3A_QA_DIR / "contact-sheet.png"
TASK_3A_PREVIEW_REPORT = TASK_3A_QA_DIR / "preview-placement.json"
TASK_3B1_QA_DIR = ROOT / "docs/publisher/qa/layout-v2/task-3b1"
TASK_3B1_RENDERER_REPORT = TASK_3B1_QA_DIR / "renderer-report.json"
TASK_3B1_CONTACT_SHEETS = (
    TASK_3B1_QA_DIR / "contact-sheet-01.png",
    TASK_3B1_QA_DIR / "contact-sheet-02.png",
    TASK_3B1_QA_DIR / "grayscale-contact-sheet-01.png",
    TASK_3B1_QA_DIR / "grayscale-contact-sheet-02.png",
)
TASK_3B1_STANDALONE_REVIEW = TASK_3B1_QA_DIR / "standalone-review.json"
NUMBERED_DIAGRAM_MANIFEST = ROOT / "docs/publisher/ru-numbered-diagrams-2026-07-15.json"
DIAGRAM_RENDERER = ROOT / "docs/publisher/tools/render_ru_inline_diagrams.mjs"
DIAGRAM_GEOMETRY_AUDIT = ROOT / "docs/publisher/tools/ru_diagram_svg_geometry.mjs"
EXPECTED_TABLE_COUNT = 12
EXPECTED_IMAGE_COUNT = 57


def test_render_qa_flags_sparse_nonblank_pages_for_review(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    for page in range(1, 4):
        (tmp_path / f"page-{page}.png").touch()

    measurements = {
        1: (0.02, (1547, 2002), (180, 180, 1360, 1900)),
        2: (0.0025, (1547, 2002), (180, 180, 1360, 1900)),
        3: (0.0004, (1547, 2002), (180, 180, 1360, 1900)),
    }

    monkeypatch.setattr(
        render_qa_metrics,
        "page_metrics",
        lambda path: measurements[render_qa_metrics.page_number(path)],
    )

    report = render_qa_metrics.collect(tmp_path, None)

    assert report["blank_like_pages"] == [3]
    assert report["sparse_pages"] == [2]
TASK_3A_REVIEW_INVENTORY_IDS = {
    "diagram-numbered-01",
    "diagram-numbered-03",
    "diagram-numbered-06",
    "diagram-inline-01",
    "diagram-inline-03",
}
TASK_3A_CHANGED_ASSET_FILENAMES = {
    "ru-figure-01-book-map.png",
    "ru-figure-03-reference-architecture.png",
    "ru-figure-04-capability-contract-path.png",
    "ru-inline-diagram-01.png",
    "ru-inline-diagram-03.png",
}
TASK_3A_EXPECTED_FILENAMES = TASK_3A_CHANGED_ASSET_FILENAMES
TASK_3B1_REVIEW_INVENTORY_IDS = {
    "diagram-numbered-02",
    "diagram-numbered-04",
    "diagram-numbered-05",
    "diagram-numbered-07",
    "diagram-numbered-08",
    "diagram-numbered-09",
    "diagram-numbered-10",
    "diagram-numbered-11",
    "diagram-numbered-12",
    "diagram-numbered-13",
    "diagram-numbered-14",
}
TASK_3B1_CHANGED_ASSET_FILENAMES = {
    "ru-figure-13-autonomy-ladder.png",
    "ru-figure-02-trust-boundaries.png",
    "ru-figure-19-localhost-control-plane.png",
    "ru-figure-16-capability-endpoint-contract.png",
    "ru-figure-06-approval-gateway.png",
    "ru-figure-25-memory-write-lifecycle.png",
    "ru-figure-05-memory-retrieval.png",
    "ru-figure-07-sandbox-mcp.png",
    "ru-figure-21-mcp-gateway.png",
    "ru-figure-08-idempotency-recovery.png",
    "ru-figure-20-eval-integrity.png",
}
TASK_3B1_EXPECTED_FILENAMES = TASK_3B1_CHANGED_ASSET_FILENAMES
ACCEPTED_TASK_3A_ASSET_HASHES = {
    "ru-figure-01-book-map.png": (
        "338e5101227b68da71fa0427c9d3dc375b975d4ed669706b124b2d83c71a6f08"
    ),
    "ru-figure-01-book-map.svg": (
        "83ed969d1fb61d97016d2c03bc3d5a351a9bb294b41a45fc5b52db7f76626f44"
    ),
    "ru-figure-03-reference-architecture.png": (
        "6d2ed8747c9b947043e91debac3c74d58b7b2c6ecb25ed4901f42bcc316300fa"
    ),
    "ru-figure-03-reference-architecture.svg": (
        "9e0e0dd78434098325bbe994f30195c7f2386d96c6baed79334cc21131ef7cb8"
    ),
    "ru-figure-04-capability-contract-path.png": (
        "fe39a37dabaa42aaedd1508cb3821e288486e1536e43e0b13ced5faabc3df610"
    ),
    "ru-figure-04-capability-contract-path.svg": (
        "da47413f2363667d4886ee6cf7fd845dcb7b18a958c1413cf935bacb1ebad631"
    ),
    "ru-inline-diagram-01.png": (
        "a857eec60ebc3e119ba812eac429b43c92625bc766b7c7937f5959a335a1c2fa"
    ),
    "ru-inline-diagram-01.svg": (
        "6b26ee456ecc274cdd5a35066626e7eb2b2cedbd6b90cf4c34bf16be4d675a43"
    ),
    "ru-inline-diagram-03.png": (
        "69e390bb6f8fb70ca42da076b202fa60d254c559584340e024f96ea3ecb25f32"
    ),
    "ru-inline-diagram-03.svg": (
        "7acd015538560e1e7bddb230304bb74d8be3f978a34e1e9af443ebdf7947187b"
    ),
}

PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
DC_NS = "http://purl.org/dc/elements/1.1/"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"


def test_publisher_layout_v2_inventory_matches_frozen_baseline() -> None:
    inventory_bytes = LAYOUT_V2_INVENTORY.read_bytes()
    inventory = json.loads(inventory_bytes)

    assert inventory_bytes == generate_publisher_layout_v2.canonical_json_bytes(inventory)
    assert inventory == generate_publisher_layout_v2.build_inventory(ROOT)
    assert inventory["schema_version"] == 2
    assert inventory["baseline_date"] == "2026-08-30"
    assert inventory["base_commit"] == "8e125f1feeb0e8ea8a61e0ef3be7e0eb3c56398a"
    assert (
        generate_publisher_layout_v2.TASK_3A_CHANGED_ASSET_FILENAMES
        == TASK_3A_CHANGED_ASSET_FILENAMES
    )
    assert (
        generate_publisher_layout_v2.TASK_3B1_CHANGED_ASSET_FILENAMES
        == TASK_3B1_CHANGED_ASSET_FILENAMES
    )
    assert inventory["source"] == {
        "path": "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
        "sha256": "8d9f8dbfccbe49efe3889cb6f67641df38f879795d2cfa2f99131fdc6ab043e8",
    }
    assert inventory["source"]["sha256"] == generate_publisher_layout_v2.sha256_text_path(
        EDITORIAL_MANUSCRIPT
    )
    assert inventory["google_doc"] == {
        "document_id": "1LyY2Psy2yaobn7VLmOwLTWm4QVrhp2I-ylE2B7V4pp4",
        "url": "https://docs.google.com/document/d/1LyY2Psy2yaobn7VLmOwLTWm4QVrhp2I-ylE2B7V4pp4/edit?tab=t.0",
        "revision": (
            "AIroW34Gzx-ZPpAtem_k8o2i17S0bE4PXlWUk3Nv_LGm1RyhlD1M4lve-"
            "opHIHp2-4Yard_T0c8j8Xs8wRWGVjaU_ninLA-x-tlweCiMzQY"
        ),
    }

    artifacts = {artifact["id"]: artifact for artifact in inventory["artifacts"]}
    assert artifacts == {
        "google-doc-book-standards": {
            "id": "google-doc-book-standards",
            "path": (
                "docs/publisher/artifacts/agent-arch-ru-google-doc-book-standards-2026-08-23.docx"
            ),
            "sha256": "3dbcbdf15be22294e7958878c78486cba7f627aaf75a8f5308d81ba5f33146a7",
            "page_count": 539,
            "page_count_report": (
                "docs/publisher/ru-google-doc-book-standards-2026-08-23.render-qa.json"
            ),
        },
        "template2000n-book-standards": {
            "id": "template2000n-book-standards",
            "path": (
                "docs/publisher/artifacts/"
                "agent-arch-ru-template2000n-book-standards-2026-08-23.docx"
            ),
            "sha256": "f9ad34503a29c105d2496af22a52365d324755d70937d1591429170499f174c3",
            "page_count": 380,
            "page_count_report": (
                "docs/publisher/ru-template2000n-book-standards-2026-08-23.render-qa.json"
            ),
        },
    }
    for artifact in artifacts.values():
        report = json.loads((ROOT / artifact["page_count_report"]).read_text())
        assert artifact["page_count"] == report["pages"]
        artifact_path = ROOT / artifact["path"]
        assert artifact["sha256"] == generate_publisher_layout_v2.sha256_binary_path(
            artifact_path
        )

    expected_counts = {
        "parts": 8,
        "chapters": 28,
        "appendices": 5,
        "formal_listings": 37,
        "fenced_code_blocks": 153,
        "manuscript_images": 57,
        "mermaid_diagrams": 56,
        "inline_diagrams": 29,
        "numbered_diagrams": 25,
        "editorial_diagrams": 2,
        "reader_facing_headings": 43,
    }
    assert inventory["counts"] == expected_counts

    diagrams = inventory["diagrams"]
    assert len(diagrams) == expected_counts["mermaid_diagrams"]
    assert len({item["id"] for item in diagrams}) == len(diagrams)
    assert len({item["filename"] for item in diagrams}) == len(diagrams)
    assert {
        family: sum(diagram["family"] == family for diagram in diagrams)
        for family in ("inline", "numbered", "editorial")
    } == {"inline": 29, "numbered": 25, "editorial": 2}
    for item in diagrams:
        assert item["node_count"] > 0
        assert item["edge_count"] > 0
        assert set(item["baseline_docx_placements"]) == set(artifacts)
        for placement in item["baseline_docx_placements"].values():
            assert placement["size_inches"]["width"] > 0
            assert placement["size_inches"]["height"] > 0

    assert len(inventory["code_blocks"]) == expected_counts["fenced_code_blocks"]
    assert len({item["id"] for item in inventory["code_blocks"]}) == len(
        inventory["code_blocks"]
    )
    assert [item["id"] for item in inventory["code_blocks"]] == [
        f"code-block-{index:03d}" for index in range(1, 154)
    ]
    assert [
        item["nearest_formal_listing"]["number"]
        for item in inventory["code_blocks"]
        if item["nearest_formal_listing"] is not None
    ] == list(range(1, expected_counts["formal_listings"] + 1))

    assert len(inventory["headings"]) == expected_counts["reader_facing_headings"]
    assert len({item["id"] for item in inventory["headings"]}) == len(inventory["headings"])
    assert {
        kind: sum(heading["kind"] == kind for heading in inventory["headings"])
        for kind in ("part", "chapter", "conclusion", "appendices", "appendix")
    } == {"part": 8, "chapter": 28, "conclusion": 1, "appendices": 1, "appendix": 5}


def test_publisher_layout_v2_review_ledger_covers_inventory_once() -> None:
    inventory = json.loads(LAYOUT_V2_INVENTORY.read_text(encoding="utf-8"))
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))

    generate_publisher_layout_v2.validate_review_ledger(inventory, ledger)
    assert ledger["schema_version"] == 2
    assert ledger["inventory_path"] == ("docs/publisher/ru-publisher-layout-v2-inventory.json")
    assert ledger["inventory_sha256"] == generate_publisher_layout_v2.sha256_json(inventory)
    assert len(ledger["entries"]) == 252
    assert len({entry["id"] for entry in ledger["entries"]}) == 252
    assert {entry["status"] for entry in ledger["entries"]} <= (
        generate_publisher_layout_v2.ALLOWED_REVIEW_STATUSES
    )


def test_final_layout_review_marks_every_diagram_as_passed() -> None:
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))
    entries = {entry["inventory_id"]: entry for entry in ledger["entries"]}

    assert all(
        entry["status"] == "pass"
        for entry in entries.values()
        if entry["item_type"] == "diagram"
    )
    assert all(
        entry["status"] == "pending"
        for entry in entries.values()
        if entry["item_type"] != "diagram"
    )

    for inventory_id in TASK_3A_REVIEW_INVENTORY_IDS:
        entry = entries[inventory_id]
        assert entry["status"] == "pass"
        assert entry["severity"] is None
        assert entry["reviewed_by"] == "Codex Task 7 final publisher review"
        assert entry["reviewed_at"] == "2026-08-30"
        assert entry["gate_statuses"] == {
            "source": "pass",
            "standalone_render": "pass",
            "preview_placement": "pass",
            "final_publisher_placement": "pass",
        }
        assert "final publisher placement pass" in entry["notes"]
        assert all(not Path(ref["path"]).is_absolute() for ref in entry["evidence_refs"])


def test_final_layout_review_closes_task_3b1_placement_gates() -> None:
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))
    entries = {entry["inventory_id"]: entry for entry in ledger["entries"]}

    for inventory_id in TASK_3B1_REVIEW_INVENTORY_IDS:
        entry = entries[inventory_id]
        assert entry["status"] == "pass"
        assert entry["severity"] is None
        assert entry["reviewed_by"] == "Codex Task 7 final publisher review"
        assert entry["reviewed_at"] == "2026-08-30"
        assert entry["gate_statuses"] == {
            "source": "pass",
            "standalone_render": "pass",
            "grayscale": "pass",
            "preview_placement": "pass",
            "final_publisher_placement": "pass",
        }
        assert "final publisher placement pass" in entry["notes"]


def test_final_layout_review_evidence_is_complete_and_current() -> None:
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))
    diagram_entries = [
        entry for entry in ledger["entries"] if entry["item_type"] == "diagram"
    ]
    required_paths = {
        str(FINAL_LAYOUT_DOCX.relative_to(ROOT)),
        str(FINAL_LAYOUT_RENDER_QA.relative_to(ROOT)),
        str(FINAL_LAYOUT_VISUAL_AUDIT.relative_to(ROOT)),
    }

    assert len(diagram_entries) == 56
    for entry in diagram_entries:
        references = {reference["path"]: reference for reference in entry["evidence_refs"]}
        assert required_paths <= set(references)
        assert any("visual-placement-contact-" in path for path in references)
        for path, reference in references.items():
            artifact = ROOT / path
            assert artifact.is_file()
            assert hashlib.sha256(artifact.read_bytes()).hexdigest() == reference["sha256"]


def test_task_3a_durable_renderer_evidence_covers_exact_assets_and_metrics() -> None:
    report = json.loads(TASK_3A_RENDERER_REPORT.read_text(encoding="utf-8"))

    assert report["mermaid"] == {
        "version": "11.17.2",
        "sha256": "581ed7d74bd9048d0e3a91363927d72ef22942d7722546b27f7cc29e35390eb8",
    }
    assert report["rendered"] == 5
    assert {result["filename"] for result in report["results"]} == (
        TASK_3A_EXPECTED_FILENAMES
    )
    assert report["minimum_effective_font_pt"] >= 9.5
    assert report["minimum_viewbox_aspect_ratio"] >= 0.72
    assert report["violations"] == []
    assert all(findings == [] for findings in report["findings"].values())
    for result in report["results"]:
        assert result["effective_font_pt"] >= 9.5
        assert result["viewbox_aspect_ratio"] >= 0.72
        assert result["violations"] == []
        assert (
            result["findings"]["viewbox_aspect_ratio"]
            == result["viewbox_aspect_ratio"]
        )
        assert all(
            findings == []
            for finding_name, findings in result["findings"].items()
            if finding_name != "viewbox_aspect_ratio"
        )

    figure_4 = next(
        result
        for result in report["results"]
        if result["filename"] == "ru-figure-04-capability-contract-path.png"
    )
    assert figure_4["effective_font_pt"] >= 9.5
    assert figure_4["viewbox_aspect_ratio"] >= 0.72


def test_task_3a_ledger_evidence_is_repository_relative_present_and_hashed() -> None:
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))
    entries = {entry["inventory_id"]: entry for entry in ledger["entries"]}

    for inventory_id in TASK_3A_REVIEW_INVENTORY_IDS:
        evidence_refs = entries[inventory_id]["evidence_refs"]
        assert {reference["kind"] for reference in evidence_refs} == {
            "renderer_report",
            "contact_sheet",
            "preview_placement",
            "svg",
            "png",
            "final_docx",
            "render_qa",
            "visual_audit",
            "final_contact_sheet",
        }
        for reference in evidence_refs:
            evidence_path = Path(reference["path"])
            assert not evidence_path.is_absolute()
            artifact_path = ROOT / evidence_path
            assert artifact_path.is_file()
            assert hashlib.sha256(artifact_path.read_bytes()).hexdigest() == reference["sha256"]

    contact_payload = TASK_3A_CONTACT_SHEET.read_bytes()
    assert contact_payload[:8] == b"\x89PNG\r\n\x1a\n"
    assert int.from_bytes(contact_payload[16:20], "big") == 2560
    assert int.from_bytes(contact_payload[20:24], "big") == 2400
    assert TASK_3A_PREVIEW_REPORT.is_file()


def test_task_3a_preview_placement_evidence_is_truthful_and_complete() -> None:
    report = json.loads(TASK_3A_PREVIEW_REPORT.read_text(encoding="utf-8"))

    assert report["schema_version"] == 1
    assert report["task"] == "3A"
    assert report["status"] == "preview_pass"
    assert report["temporary_docx_committed"] is False
    assert report["final_publisher_placement"] == "pending_task_7"
    assert report["visual_review"] == {
        "result": "pass",
        "reviewed_by": "Codex",
        "reviewed_on": "2026-08-30",
    }

    placements = report["placements"]
    assert {placement["filename"] for placement in placements} == (
        TASK_3A_EXPECTED_FILENAMES
    )
    assert {placement["page_number"] for placement in placements} <= set(
        report["rendered_pages"]
    )
    assert len(report["rendered_pages"]) == len(set(report["rendered_pages"]))
    for placement in placements:
        asset_path = ROOT / "docs/publisher/visuals" / placement["filename"]
        assert placement["asset_sha256"] == hashlib.sha256(asset_path.read_bytes()).hexdigest()
        assert placement["payload_match"] is True
        assert placement["size_inches"]["width"] > 0
        assert placement["size_inches"]["height"] > 0
        assert placement["effective_font_pt"] >= 9.5
        assert placement["viewbox_aspect_ratio"] >= 0.72
        assert placement["visually_verified"] is True


def test_task_3b1_durable_renderer_evidence_covers_exact_assets_and_metrics() -> None:
    report = json.loads(TASK_3B1_RENDERER_REPORT.read_text(encoding="utf-8"))

    assert report["mermaid"] == {
        "version": "11.17.2",
        "sha256": "581ed7d74bd9048d0e3a91363927d72ef22942d7722546b27f7cc29e35390eb8",
    }
    assert report["rendered"] == 11
    assert {result["filename"] for result in report["results"]} == (
        TASK_3B1_EXPECTED_FILENAMES
    )
    assert report["minimum_effective_font_pt"] >= 9.5
    assert report["minimum_viewbox_aspect_ratio"] == min(
        result["viewbox_aspect_ratio"] for result in report["results"]
    )
    assert report["violations"] == []
    expected_overrides = generate_publisher_layout_v2.TASK_3B1_ASPECT_RATIO_OVERRIDES
    assert {
        finding["filename"]: finding["review"]
        for finding in report["findings"]["aspect_ratio_overrides"]
    } == expected_overrides
    assert all(
        findings == []
        for finding_name, findings in report["findings"].items()
        if finding_name != "aspect_ratio_overrides"
    )
    assert {result["layout_engine"] for result in report["results"]} == {"dagre"}
    assert {result["layout_class"] for result in report["results"]} == {
        "simple-flow",
        "decision-state",
        "evidence-overlay",
    }
    for result in report["results"]:
        assert result["source_sha256"]
        assert result["svg_sha256"]
        assert result["png_sha256"]
        assert result["effective_font_pt"] >= 9.5
        expected_override = expected_overrides.get(result["filename"])
        if expected_override is None:
            assert result["viewbox_aspect_ratio"] >= 0.72
            assert result["aspect_ratio_override"] is None
        else:
            assert result["viewbox_aspect_ratio"] < 0.72
            assert result["aspect_ratio_override"] == expected_override
        assert result["violations"] == []
        assert all(
            findings == []
            for finding_name, findings in result["findings"].items()
            if finding_name not in {"viewbox_aspect_ratio", "aspect_ratio_overrides"}
        )

    source_by_filename = {
        item["filename"]: item["mermaid"]
        for item in json.loads(NUMBERED_DIAGRAM_MANIFEST.read_text(encoding="utf-8"))[
            "diagrams"
        ]
    }
    for result in report["results"]:
        source = source_by_filename[result["filename"]]
        assert result["source_sha256"] == (
            generate_publisher_layout_v2.mermaid_source_sha256(source)
        )
        generate_publisher_layout_v2.validate_rendered_diagram_evidence(
            ROOT,
            result,
            source,
        )


@pytest.mark.parametrize(
    ("mutation", "message"),
    [
        ("source", "source SHA-256"),
        ("reported_source_hash", "source SHA-256"),
        ("svg", "SVG SHA-256"),
        ("png", "PNG SHA-256"),
    ],
)
def test_task_3b1_validator_rejects_stale_source_or_assets(
    tmp_path: Path,
    mutation: str,
    message: str,
) -> None:
    renderer_report = json.loads(TASK_3B1_RENDERER_REPORT.read_text(encoding="utf-8"))
    result = copy.deepcopy(renderer_report["results"][0])
    source_by_filename = {
        item["filename"]: item["mermaid"]
        for item in json.loads(NUMBERED_DIAGRAM_MANIFEST.read_text(encoding="utf-8"))[
            "diagrams"
        ]
    }
    source = source_by_filename[result["filename"]]
    output_dir = tmp_path / "docs/publisher/visuals"
    output_dir.mkdir(parents=True)
    svg_source = ROOT / "docs/publisher/visuals" / result["svg"]
    png_source = ROOT / "docs/publisher/visuals" / result["png"]
    (output_dir / result["svg"]).write_bytes(svg_source.read_bytes())
    (output_dir / result["png"]).write_bytes(png_source.read_bytes())

    if mutation == "source":
        source = f"{source}\n"
    elif mutation == "reported_source_hash":
        result["source_sha256"] = "0" * 64
    elif mutation == "svg":
        (output_dir / result["svg"]).write_bytes(svg_source.read_bytes() + b"\n")
    elif mutation == "png":
        (output_dir / result["png"]).write_bytes(png_source.read_bytes() + b"stale")

    with pytest.raises(ValueError, match=message):
        generate_publisher_layout_v2.validate_rendered_diagram_evidence(
            tmp_path,
            result,
            source,
        )


def test_task_3b1_evidence_is_repository_relative_present_and_hashed() -> None:
    ledger = json.loads(LAYOUT_V2_LEDGER.read_text(encoding="utf-8"))
    entries = {entry["inventory_id"]: entry for entry in ledger["entries"]}

    expected_common_paths = {
        path.relative_to(ROOT).as_posix()
        for path in (
            TASK_3B1_RENDERER_REPORT,
            *TASK_3B1_CONTACT_SHEETS,
            TASK_3B1_STANDALONE_REVIEW,
        )
    }
    for inventory_id in TASK_3B1_REVIEW_INVENTORY_IDS:
        evidence_refs = entries[inventory_id]["evidence_refs"]
        evidence_paths = {reference["path"] for reference in evidence_refs}
        assert expected_common_paths <= evidence_paths
        assert len(evidence_refs) == 12
        for reference in evidence_refs:
            evidence_path = Path(reference["path"])
            assert not evidence_path.is_absolute()
            artifact_path = ROOT / evidence_path
            assert artifact_path.is_file()
            assert hashlib.sha256(artifact_path.read_bytes()).hexdigest() == reference["sha256"]

    for contact_sheet in TASK_3B1_CONTACT_SHEETS:
        payload = contact_sheet.read_bytes()
        assert payload[:8] == b"\x89PNG\r\n\x1a\n"
        assert int.from_bytes(payload[16:20], "big") == 4200
        assert int.from_bytes(payload[20:24], "big") == 3900


def test_task_3b1_review_records_only_standalone_gates_as_passed() -> None:
    report = json.loads(TASK_3B1_STANDALONE_REVIEW.read_text(encoding="utf-8"))

    assert report["schema_version"] == 3
    assert report["task"] == "3B1"
    assert report["status"] == "standalone_pass"
    assert report["temporary_docx_committed"] is False
    assert report["page_renders_committed"] is False
    assert report["preview_placement"] == "pending_task_7"
    assert report["final_publisher_placement"] == "pending_task_7"
    assert set(report["selected_assets"]) == TASK_3B1_EXPECTED_FILENAMES
    assert report["visual_review"] == {
        "result": "pass",
        "reviewed_by": "Kant (independent layout review)",
        "reviewed_on": "2026-08-30",
        "standalone_render": "pass",
        "grayscale": "pass",
    }

    assert "rendered_pages" not in report
    assert "placements" not in report
    asset_reviews = report["asset_reviews"]
    assert {review["filename"] for review in asset_reviews} == TASK_3B1_EXPECTED_FILENAMES
    for review in asset_reviews:
        assert review["gates"] == {
            "source": "pass",
            "standalone_render": "pass",
            "grayscale": "pass",
            "preview_placement": "pending",
            "final_publisher_placement": "pending",
        }
        assert review["source_sha256"]
        assert review["svg_sha256"]
        assert review["png_sha256"]


def test_task_3b1_does_not_change_the_five_accepted_diagram_assets() -> None:
    for filename, expected_hash in ACCEPTED_TASK_3A_ASSET_HASHES.items():
        asset_path = ROOT / "docs/publisher/visuals" / filename
        assert hashlib.sha256(asset_path.read_bytes()).hexdigest() == expected_hash


def paragraph_uses_monospace_font(paragraph: ET.Element) -> bool:
    styled_characters = 0
    monospace_characters = 0
    monospace_run_seen = False
    for run in paragraph.findall(f"{{{WORD_NS}}}r"):
        run_text = "".join(
            node.text or "" for node in run.findall(f".//{{{WORD_NS}}}t")
        )
        fonts = run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}rFonts")
        if fonts is None:
            continue
        is_monospace = any(
            "mono" in value.lower() or "courier" in value.lower()
            for value in fonts.attrib.values()
        )
        monospace_run_seen = monospace_run_seen or is_monospace
        if not run_text.strip():
            continue
        styled_characters += len(run_text)
        if is_monospace:
            monospace_characters += len(run_text)
    if styled_characters == 0:
        return monospace_run_seen
    return monospace_characters / styled_characters >= 0.8


def active_on_off_property(paragraph: ET.Element, name: str) -> bool:
    property_node = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}{name}")
    return property_node is not None and property_node.attrib.get(
        f"{{{WORD_NS}}}val", "1"
    ) not in {"0", "false", "off"}


def test_visual_audit_accepts_every_current_manuscript_image() -> None:
    assert (
        len(sync_ru_docx_visuals.parse_manuscript_visuals(EDITORIAL_MANUSCRIPT))
        == EXPECTED_IMAGE_COUNT
    )


def test_mermaid_visuals_reserve_space_below_cluster_titles() -> None:
    report = json.loads(VISUAL_LAYOUT_AUDIT.read_text(encoding="utf-8"))

    assert report["rendered"] == 56
    assert report["violations"] == []
    assert report["minimum_effective_font_pt"] >= 8.5
    assert report["cluster_title_violations"] == []
    assert report["cluster_title_edge_violations"] == []
    assert report["node_label_violations"] == []
    assert report["minimum_cluster_title_gap_px"] >= 12


def test_mermaid_renderer_rejects_hidden_edges_and_intrinsic_label_clipping() -> None:
    source = DIAGRAM_RENDERER.read_text(encoding="utf-8") + (
        DIAGRAM_GEOMETRY_AUDIT.read_text(encoding="utf-8")
    )

    assert "if (crossedLabel.opaque_background) continue;" not in source
    assert "scrollWidth" in source
    assert "clientWidth" in source


def test_generated_mermaid_svgs_are_valid_standalone_xml() -> None:
    for manifest_path in (INLINE_DIAGRAMS, NUMBERED_DIAGRAMS):
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        for diagram in manifest["diagrams"]:
            svg_path = ROOT / "docs/publisher/visuals" / diagram["filename"].replace(
                ".png", ".svg"
            )
            ET.parse(svg_path)


def test_only_figure_3_uses_reviewed_lane_frame_transitions() -> None:
    for manifest_path in (INLINE_DIAGRAMS, NUMBERED_DIAGRAMS):
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        for diagram in manifest["diagrams"]:
            source = diagram["mermaid"]
            cluster_ids = set(re.findall(r"^subgraph\s+([A-Za-z][\w-]*)", source, re.M))
            linked_cluster_ids = {
                endpoint
                for left, right in re.findall(
                    r"^\s*([A-Za-z][\w-]*)\s*(?:-->|-\.->|~~~)(?:\|[^\n]*?\|)?\s*"
                    r"([A-Za-z][\w-]*)",
                    source,
                    re.M,
                )
                for endpoint in (left, right)
                if endpoint in cluster_ids
            }
            expected_cluster_ids = (
                {"REQUEST", "CONTROL", "EVIDENCE"}
                if diagram["filename"] == "ru-figure-03-reference-architecture.png"
                else set()
            )
            assert linked_cluster_ids == expected_cluster_ids, (
                f"{diagram['filename']} links cluster frames: "
                f"{sorted(linked_cluster_ids)}"
            )


def test_independently_reviewed_decisions_have_explicit_safe_branches() -> None:
    inline = {
        item["filename"]: item["mermaid"]
        for item in json.loads(INLINE_DIAGRAMS.read_text(encoding="utf-8"))["diagrams"]
    }

    assert 'V -->|"Да"| Y' in inline["ru-inline-diagram-05.png"]
    assert 'V -->|"Нет"| Q' in inline["ru-inline-diagram-05.png"]
    assert 'U -->|"Нет"| E' in inline["ru-inline-diagram-10.png"]
    assert 'U -->|"Да"| F' in inline["ru-inline-diagram-10.png"]
    assert 'F -->|"Эффекта нет"| E' in inline["ru-inline-diagram-10.png"]
    assert 'F -->|"Исход неизвестен"| H' in inline["ru-inline-diagram-10.png"]
    assert 'G -->|"Да"| H' in inline["ru-inline-diagram-29.png"]
    assert 'G -->|"Нет"| R' in inline["ru-inline-diagram-29.png"]


def test_visual_audit_uses_the_parsed_manuscript_count_for_docx_validation() -> None:
    validate = getattr(sync_ru_docx_visuals, "validate_docx_image_counts", None)

    assert validate is not None
    validate(EXPECTED_IMAGE_COUNT, EXPECTED_IMAGE_COUNT, EXPECTED_IMAGE_COUNT)
    with pytest.raises(ValueError, match="raw=56, template=57, expected=57"):
        validate(56, 57, 57)


def test_visual_resize_is_idempotent_at_the_print_height_limit() -> None:
    drawing = ET.fromstring(
        f'<w:drawing xmlns:w="{WORD_NS}" xmlns:wp="{DRAWING_NS}" '
        f'xmlns:a="{DRAWINGML_NS}">'
        '<wp:extent cx="5943600" cy="5943600"/>'
        '<a:xfrm><a:ext cx="5943600" cy="5943600"/></a:xfrm>'
        "</w:drawing>"
    )
    image = ROOT / "docs/publisher/visuals/ru-figure-13-autonomy-ladder.png"

    first = sync_ru_docx_visuals.resize_drawing(drawing, image)
    first_xml = ET.tostring(drawing)
    second = sync_ru_docx_visuals.resize_drawing(drawing, image)

    assert first == second
    assert first[1] <= sync_ru_docx_visuals.MAX_FIGURE_HEIGHT
    assert ET.tostring(drawing) == first_xml


def test_visual_resize_expands_small_legacy_placeholders_to_the_print_frame() -> None:
    drawing = ET.fromstring(
        f'<w:drawing xmlns:w="{WORD_NS}" xmlns:wp="{DRAWING_NS}" '
        f'xmlns:a="{DRAWINGML_NS}">'
        f'<wp:extent cx="{sync_ru_docx_visuals.EMU_PER_INCH}" '
        f'cy="{sync_ru_docx_visuals.EMU_PER_INCH}"/>'
        f'<a:xfrm><a:ext cx="{sync_ru_docx_visuals.EMU_PER_INCH}" '
        f'cy="{sync_ru_docx_visuals.EMU_PER_INCH}"/></a:xfrm>'
        "</w:drawing>"
    )
    image = ROOT / "docs/publisher/visuals/ru-figure-01-book-map.png"

    width, height = sync_ru_docx_visuals.resize_drawing(drawing, image)

    assert width > sync_ru_docx_visuals.EMU_PER_INCH
    assert width <= sync_ru_docx_visuals.MAX_FIGURE_WIDTH
    assert height <= sync_ru_docx_visuals.MAX_FIGURE_HEIGHT


def test_caption_normalizer_moves_only_captions_that_precede_images() -> None:
    paragraphs: list[str] = []
    for number in range(1, 26):
        caption = f"<w:p><w:r><w:t>Рисунок {number}. Схема {number}</w:t></w:r></w:p>"
        image = "<w:p><w:r><w:drawing/></w:r></w:p>"
        paragraphs.extend((caption, image) if number <= 20 else (image, caption))
    payload = (
        f'<w:document xmlns:w="{WORD_NS}"><w:body>'
        + "".join(paragraphs)
        + "</w:body></w:document>"
    ).encode()

    normalized, metrics = normalize_docx_figure_caption_order.normalize_document_xml(
        payload
    )
    document = ET.fromstring(normalized)
    children = list(document.find(f"{{{WORD_NS}}}body"))

    assert metrics == {
        "numbered_captions": 25,
        "captions_moved": 20,
        "captions_already_after_images": 5,
        "captions_after_images": 25,
    }
    for index, paragraph in enumerate(children):
        text = "".join(paragraph.itertext()).strip()
        if not text.startswith("Рисунок "):
            continue
        assert index > 0
        assert children[index - 1].find(f".//{{{WORD_NS}}}drawing") is not None


def test_editorial_builder_knows_all_eight_part_boundaries() -> None:
    tree = ast.parse(EDITORIAL_BUILDER.read_text(encoding="utf-8"))
    assignment = next(
        node
        for node in tree.body
        if isinstance(node, ast.Assign)
        and any(
            isinstance(target, ast.Name) and target.id == "FIRST_CHAPTERS_IN_PART"
            for target in node.targets
        )
    )
    assert ast.literal_eval(assignment.value) == {1, 4, 7, 10, 13, 17, 22, 26}


def test_editorial_renderer_adds_resolvable_internal_bookmarks_and_links(
    tmp_path: Path,
) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "internal-links.docx"
    script = r'''
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    """
    <h2>Глава 1. Начало</h2>
    <p>Продолжение находится в главе 16; см. рисунок 3, таблицу 4 и листинг 12.</p>
    <h2>Глава 16. Доказательства</h2>
    <p>Рисунок 3. Цепочка доказательств</p>
    <p>Таблица 4. Выпускные сигналы</p>
    <p><strong>Листинг 12. Проверка решения.</strong></p>
    """,
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
'''
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    bookmarks = document.findall(f".//{{{WORD_NS}}}bookmarkStart")
    names = [node.attrib[f"{{{WORD_NS}}}name"] for node in bookmarks]
    identifiers = [node.attrib[f"{{{WORD_NS}}}id"] for node in bookmarks]
    assert names == ["ch_1", "ch_16", "fig_3", "table_4", "listing_12"]
    assert len(identifiers) == len(set(identifiers))

    hyperlinks = document.findall(f".//{{{WORD_NS}}}hyperlink")
    anchors = [
        node.attrib[f"{{{WORD_NS}}}anchor"]
        for node in hyperlinks
        if f"{{{WORD_NS}}}anchor" in node.attrib
    ]
    assert anchors == ["ch_16", "fig_3", "table_4", "listing_12"]
    assert set(anchors) <= set(names)

    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        value = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
        if re.match(r"^(?:Рисунок 3|Таблица 4|Листинг 12)\.", value):
            assert paragraph.find(f"{{{WORD_NS}}}hyperlink") is None


def test_editorial_renderer_compacts_sources_into_a_breakable_run_in_bibliography(
    tmp_path: Path,
) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "source-paragraph.docx"
    script = r'''
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    """
    <h3>Источники главы</h3>
    <p><strong>S001.</strong> OWASP, AI Agent Security Cheat Sheet.</p>
    <p><strong>S002.</strong> NIST, AI RMF 1.0.</p>
    """,
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
'''
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    source_paragraphs = []
    source_heading = None
    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        value = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
        if value.startswith("Источники главы"):
            source_heading = paragraph
        if re.match(r"^S\d{3}\.", value):
            source_paragraphs.append(paragraph)

    assert source_heading is not None
    heading_spacing = source_heading.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
    assert heading_spacing is not None
    assert int(heading_spacing.attrib[f"{{{WORD_NS}}}before"]) == 0
    assert int(heading_spacing.attrib[f"{{{WORD_NS}}}after"]) == 0
    assert int(heading_spacing.attrib[f"{{{WORD_NS}}}line"]) == 210
    assert heading_spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "exact"
    heading_sizes = {
        int(size.attrib[f"{{{WORD_NS}}}val"])
        for size in source_heading.findall(f".//{{{WORD_NS}}}rPr/{{{WORD_NS}}}sz")
    }
    assert heading_sizes == {16, 20}
    assert "".join(source_heading.itertext()) == "Источники главы: S001, S002"
    assert source_heading.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is None
    assert source_paragraphs == []


def test_editorial_renderer_formats_source_lists_as_compact_breakable_items(
    tmp_path: Path,
) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "source-list.docx"
    script = r'''
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    """
    <h3>Источники главы</h3>
    <ul>
      <li><strong>S001.</strong> OWASP, AI Agent Security Cheat Sheet.</li>
      <li><strong>S002.</strong> NIST, AI RMF 1.0.</li>
    </ul>
    """,
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
'''
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    source_items = []
    source_heading = None
    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        value = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
        if value.startswith("Источники главы"):
            source_heading = paragraph
        if re.match(r"^S\d{3}\.", value):
            source_items.append(paragraph)

    assert source_heading is not None
    assert "".join(source_heading.itertext()) == "Источники главы: S001, S002"
    assert source_items == []


def test_editorial_renderer_uses_compact_spacing_for_regular_lists(
    tmp_path: Path,
) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "regular-list.docx"
    script = r'''
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    "<ul><li>Первый пункт</li><li>Второй пункт</li></ul>",
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
'''
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    list_items = document.findall(f".//{{{WORD_NS}}}p")
    assert len(list_items) == 2
    for paragraph in list_items:
        spacing = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
        assert spacing is not None
        assert int(spacing.attrib[f"{{{WORD_NS}}}before"]) == 0
        assert int(spacing.attrib[f"{{{WORD_NS}}}after"]) == 40
        assert int(spacing.attrib[f"{{{WORD_NS}}}line"]) == 260
        assert spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "atLeast"
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepLines") is not None


def test_editorial_document_enables_widow_control_for_body_paragraphs(
    tmp_path: Path,
) -> None:
    document = build_ru_editorial_docx.Document()

    build_ru_editorial_docx.configure_document_properties(document)
    output = tmp_path / "widow-control.docx"
    document.save(output)

    with ZipFile(output) as archive:
        styles = ET.fromstring(archive.read("word/styles.xml"))

    normal = next(
        style
        for style in styles.findall(f".//{{{WORD_NS}}}style")
        if style.attrib.get(f"{{{WORD_NS}}}styleId") == "Normal"
    )
    widow_control = normal.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}widowControl")
    assert widow_control is not None
    assert widow_control.attrib.get(f"{{{WORD_NS}}}val", "true") not in {
        "false",
        "0",
        "off",
    }


def test_editorial_renderer_compacts_lab_step_headings(tmp_path: Path) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "lab-step-heading.docx"
    script = r'''
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    """
    <h4>Шаг 1. Зафиксируйте сценарий</h4>
    <p>Опишите вход, результат и владельца последствия.</p>
    """,
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
'''
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    heading = next(
        paragraph
        for paragraph in document.findall(f".//{{{WORD_NS}}}p")
        if "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        ).startswith("Шаг 1.")
    )
    spacing = heading.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
    assert spacing is not None
    assert int(spacing.attrib[f"{{{WORD_NS}}}before"]) <= 80
    assert int(spacing.attrib[f"{{{WORD_NS}}}after"]) <= 40
    assert active_on_off_property(heading, "keepNext")


def test_editorial_renderer_compacts_key_takeaway_blocks(tmp_path: Path) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "key-takeaways.docx"
    document = Document()
    root = lxml_html.fragment_fromstring(
        """
        <h3>Ключевые выводы</h3>
        <ul>
          <li>Первый вывод.</li>
          <li>Второй вывод.</li>
          <li>Третий вывод.</li>
        </ul>
        """,
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    paragraphs = document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
    heading = next(
        paragraph
        for paragraph in paragraphs
        if "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        )
        == "Ключевые выводы"
    )
    heading_spacing = heading.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
    assert heading_spacing is not None
    assert int(heading_spacing.attrib[f"{{{WORD_NS}}}before"]) <= 120
    assert int(heading_spacing.attrib[f"{{{WORD_NS}}}after"]) <= 40
    assert {
        int(size.attrib[f"{{{WORD_NS}}}val"])
        for size in heading.findall(f".//{{{WORD_NS}}}rPr/{{{WORD_NS}}}sz")
    } == {26}

    takeaway_items = [
        paragraph
        for paragraph in paragraphs
        if "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        ).endswith("вывод.")
    ]
    assert len(takeaway_items) == 3
    for paragraph in takeaway_items:
        spacing = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
        assert spacing is not None
        assert int(spacing.attrib[f"{{{WORD_NS}}}before"]) == 0
        assert int(spacing.attrib[f"{{{WORD_NS}}}after"]) == 0
        assert int(spacing.attrib[f"{{{WORD_NS}}}line"]) == 210
        assert spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "exact"
        assert {
            int(size.attrib[f"{{{WORD_NS}}}val"])
            for size in paragraph.findall(f".//{{{WORD_NS}}}rPr/{{{WORD_NS}}}sz")
        } == {19}


def test_editorial_renderer_compacts_practical_step_paragraph(tmp_path: Path) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "practical-step.docx"
    document = Document()
    root = lxml_html.fragment_fromstring(
        """
        <h2>Глава 1. Проверка</h2>
        <p><strong>Практический шаг.</strong>
        Сопоставьте решение с фактическими вызовами и владельцем.</p>
        """,
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    paragraph = next(
        paragraph
        for paragraph in document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
        if "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        ).startswith("Практический шаг.")
    )
    spacing = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
    assert spacing is not None
    assert int(spacing.attrib[f"{{{WORD_NS}}}before"]) == 0
    assert int(spacing.attrib[f"{{{WORD_NS}}}after"]) == 0
    assert int(spacing.attrib[f"{{{WORD_NS}}}line"]) == 210
    assert spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "exact"
    assert active_on_off_property(paragraph, "keepLines")


def test_editorial_renderer_compacts_lab_support_paragraphs(tmp_path: Path) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "lab-support-paragraphs.docx"
    document = Document()
    root = lxml_html.fragment_fromstring(
        """
        <h2>Глава 1. Проверка</h2>
        <p><strong>Наблюдение.</strong> Проверьте связность трассы.</p>
        <p><strong>Критерий приемки:</strong> артефакты согласованы.</p>
        <p><strong>Что доказывает результат.</strong> Связь воспроизводима.</p>
        <p><strong>Накопительный артефакт.</strong> Сохраните результат.</p>
        <p><strong>Отрицательная проверка.</strong> Проверьте отказ.</p>
        <p><strong>Если результат отличается.</strong> Зафиксируйте расхождение.</p>
        <p><strong>Дополнительное задание.</strong> Проверьте резервный маршрут.</p>
        """,
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    paragraphs = [
        paragraph
        for paragraph in document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
        if not "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        ).startswith("Глава 1.")
    ]
    assert len(paragraphs) == 7
    for paragraph in paragraphs:
        spacing = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
        assert spacing is not None
        assert int(spacing.attrib[f"{{{WORD_NS}}}before"]) == 0
        assert int(spacing.attrib[f"{{{WORD_NS}}}after"]) == 0
        assert int(spacing.attrib[f"{{{WORD_NS}}}line"]) == 220
        assert spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "exact"
        assert {
            int(size.attrib[f"{{{WORD_NS}}}val"])
            for size in paragraph.findall(f".//{{{WORD_NS}}}rPr/{{{WORD_NS}}}sz")
        } == {20}


def test_editorial_renderer_keeps_multiline_code_blocks_together(tmp_path: Path) -> None:
    runtime_python = Path(
        os.environ.get(
            "CODEX_DOCUMENT_PYTHON",
            Path.home()
            / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3",
        )
    )
    if not runtime_python.is_file():
        pytest.skip("bundled document runtime is unavailable")

    output = tmp_path / "code-block.docx"
    script = r"""
import sys
from pathlib import Path
from docx import Document
from lxml import html

sys.path.insert(0, sys.argv[1])
from docs.publisher.tools import build_ru_editorial_docx

document = Document()
root = html.fragment_fromstring(
    "<h2>Глава 1. Тест</h2>"
    "<p><strong>Команда.</strong> Запустите сценарий:</p>"
    "<pre><code>first line\nsecond line\nlast line</code></pre>",
    create_parent="div",
)
renderer = build_ru_editorial_docx.DocxRenderer(
    document,
    Path(sys.argv[1]) / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
)
renderer.render(root)
document.save(sys.argv[2])
"""
    subprocess.run(
        [str(runtime_python), "-c", script, str(ROOT), str(output)],
        cwd=ROOT,
        check=True,
    )

    with ZipFile(output) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    intro_paragraph = None
    code_paragraphs = []
    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        value = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
        if value == "Команда. Запустите сценарий:":
            intro_paragraph = paragraph
        elif value in {"first line", "second line", "last line"}:
            code_paragraphs.append(paragraph)

    assert intro_paragraph is not None
    assert intro_paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
    assert len(code_paragraphs) == 3
    for paragraph in code_paragraphs:
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepLines") is not None
    for paragraph in code_paragraphs[:-1]:
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
    assert code_paragraphs[-1].find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is None


def test_editorial_renderer_keeps_page_sized_formal_listing_together(
    tmp_path: Path,
) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "page-sized-listing.docx"
    code = "\n".join(f"value_{index} = {index}" for index in range(1, 25))
    document = Document()
    root = lxml_html.fragment_fromstring(
        (
            "<h2>Глава 1. Проверка листинга</h2>"
            "<p><strong>Листинг 1. Проверяемая функция.</strong> Тип: учебный пример.</p>"
            "<p><strong>Как читать листинг.</strong> Проследите все строки.</p>"
            f'<pre><code class="language-python">{code}</code></pre>'
        ),
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    reading_paragraph = None
    code_paragraphs = []
    for paragraph in document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p"):
        value = "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        )
        if value.startswith("Как читать листинг."):
            reading_paragraph = paragraph
        elif re.match(r"^\s*\d+\s+value_\d+", value):
            code_paragraphs.append(paragraph)

    assert reading_paragraph is not None
    assert reading_paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
    assert len(code_paragraphs) == 24
    for paragraph in code_paragraphs[:-1]:
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
    assert code_paragraphs[-1].find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is None


def test_editorial_renderer_keeps_short_text_output_together(tmp_path: Path) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "text-output.docx"
    code = "\n".join(f"artifact-{index:02d}" for index in range(1, 19))
    document = Document()
    root = lxml_html.fragment_fromstring(
        f'<pre><code class="language-text">{code}</code></pre>',
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    paragraphs = document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
    assert len(paragraphs) == 18
    for paragraph in paragraphs[:-1]:
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
    assert paragraphs[-1].find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is None


def test_editorial_renderer_highlights_code_and_numbers_only_formal_listings(
    tmp_path: Path,
) -> None:
    from docx import Document
    from lxml import html as lxml_html

    from docs.publisher.tools import build_ru_editorial_docx

    output = tmp_path / "highlighted-listing.docx"
    document = Document()
    root = lxml_html.fragment_fromstring(
        """
        <h2>Глава 1. Проверка листинга</h2>
        <p><strong>Листинг 1. Проверяемая функция.</strong> Тип: учебный пример.</p>
        <p><strong>Как читать листинг.</strong> Проследите ветвление.</p>
        <pre><code class="language-python">def greet(name):
    # Комментарий
    return &quot;ok&quot;</code></pre>
        <p>Команда проверки:</p>
        <pre><code class="language-console">python -m pytest</code></pre>
        """,
        create_parent="div",
    )
    renderer = build_ru_editorial_docx.DocxRenderer(
        document,
        ROOT / "docs/publisher/ru-manuscript-editorial-2026-07-13.md",
    )
    renderer.render(root)
    document.save(output)

    with ZipFile(output) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))

    paragraphs = document_xml.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
    code_blocks: list[list[ET.Element]] = []
    current: list[ET.Element] = []
    for paragraph in paragraphs:
        if paragraph_uses_monospace_font(paragraph):
            current.append(paragraph)
        elif current:
            code_blocks.append(current)
            current = []
    if current:
        code_blocks.append(current)

    assert len(code_blocks) == 2
    formal, command = code_blocks
    assert [
        "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
        for paragraph in formal
    ] == [
        " 1  def greet(name):",
        " 2      # Комментарий",
        ' 3      return "ok"',
    ]
    assert "".join(
        node.text or "" for node in command[0].findall(f".//{{{WORD_NS}}}t")
    ) == "python -m pytest"

    for paragraph in formal + command:
        spacing = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}spacing")
        assert spacing is not None
        assert spacing.attrib[f"{{{WORD_NS}}}lineRule"] == "exact"
        assert spacing.attrib[f"{{{WORD_NS}}}line"] == "210"
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}suppressAutoHyphens") is not None
        assert all(
            run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}noProof") is not None
            for run in paragraph.findall(f"{{{WORD_NS}}}r")
        )
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pBdr/{{{WORD_NS}}}left") is not None

    assert formal[0].find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pBdr/{{{WORD_NS}}}top") is not None
    assert formal[-1].find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pBdr/{{{WORD_NS}}}bottom") is not None
    colors = {
        color.attrib[f"{{{WORD_NS}}}val"]
        for paragraph in formal
        for color in paragraph.findall(f"{{{WORD_NS}}}r/{{{WORD_NS}}}rPr/{{{WORD_NS}}}color")
    }
    assert {"234E8A", "667085", "276749", "7A8088"} <= colors

    listing_caption = next(
        paragraph
        for paragraph in paragraphs
        if "".join(
            node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
        ).startswith("Листинг 1.")
    )
    assert active_on_off_property(listing_caption, "keepNext")
    assert renderer.metrics["formal_listings"] == 1
    assert renderer.metrics["numbered_code_lines"] == 3
    assert renderer.metrics["code_lines"] == 4
    assert renderer.metrics["highlighted_code_runs"] >= 4


def test_syntax_highlighting_dependency_cannot_fail_silently(monkeypatch) -> None:
    import builtins

    from docs.publisher.tools import build_ru_editorial_docx

    original_import = builtins.__import__

    def import_without_pygments(name, *args, **kwargs):
        if name == "pygments" or name.startswith("pygments."):
            raise ImportError("simulated missing Pygments")
        return original_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", import_without_pygments)

    with pytest.raises(RuntimeError, match="Pygments is required"):
        build_ru_editorial_docx.highlighted_code_lines("def greet():\n    return 1", "python")

    assert build_ru_editorial_docx.highlighted_code_lines("plain text", "text") == [
        [("plain text", build_ru_editorial_docx.CodeRunStyle())]
    ]


def ordered_embedded_images(path: Path) -> tuple[list[str], list[str]]:
    with ZipFile(path) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        relationships = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        targets = {
            node.attrib["Id"]: posixpath.normpath(f"word/{node.attrib['Target']}")
            for node in relationships.findall(f"{{{PACKAGE_REL_NS}}}Relationship")
            if node.attrib.get("Type", "").endswith("/image")
        }
        ordered_targets = [
            targets[blip.attrib[f"{{{OFFICE_REL_NS}}}embed"]]
            for blip in document.findall(f".//{{{DRAWINGML_NS}}}blip")
        ]
        hashes = [hashlib.sha256(archive.read(target)).hexdigest() for target in ordered_targets]
    return ordered_targets, hashes


def test_template2000n_final_preserves_embedded_font_registrations() -> None:
    with ZipFile(EDITORIAL_TEMPLATE_DOCX) as archive:
        names = set(archive.namelist())
        font_table = ET.fromstring(archive.read("word/fontTable.xml"))
        relationships = ET.fromstring(archive.read("word/_rels/fontTable.xml.rels"))
        numbering_xml = archive.read("word/numbering.xml")

    font_relationships = {
        node.attrib["Id"]: node.attrib["Target"]
        for node in relationships.findall(f"{{{PACKAGE_REL_NS}}}Relationship")
        if node.attrib.get("Type", "").endswith("/font")
    }
    registered_relationships = {
        node.attrib[f"{{{OFFICE_REL_NS}}}id"]
        for tag in ("embedRegular", "embedBold", "embedItalic", "embedBoldItalic")
        for node in font_table.findall(f".//{{{WORD_NS}}}{tag}")
    }

    assert font_relationships
    assert font_relationships.keys() == registered_relationships
    assert {f"word/{target}" for target in font_relationships.values()} <= names
    font_names = {
        node.attrib[f"{{{WORD_NS}}}name"] for node in font_table.findall(f"{{{WORD_NS}}}font")
    }
    assert "Roboto Mono" in font_names
    assert b"Noto Sans Symbols" not in numbering_xml


def test_current_book_standards_artifacts_pass_machine_gates() -> None:
    raw_a11y = json.loads(RAW_A11Y.read_text(encoding="utf-8"))
    raw_render = json.loads(RAW_RENDER_QA.read_text(encoding="utf-8"))
    template_a11y = json.loads(TEMPLATE_A11Y.read_text(encoding="utf-8"))
    font_audit = json.loads(TEMPLATE_FONT_AUDIT.read_text(encoding="utf-8"))
    metrics = json.loads(TEMPLATE_METRICS.read_text(encoding="utf-8"))
    template_render = json.loads(TEMPLATE_RENDER_QA.read_text(encoding="utf-8"))
    visual_audit = json.loads(TEMPLATE_VISUAL_AUDIT.read_text(encoding="utf-8"))

    assert raw_a11y["counts"] == {"high": 0, "medium": 0, "low": 2}
    assert template_a11y["counts"] == {"high": 0, "medium": 0, "low": 2}
    assert raw_render["pages"] == 539
    assert raw_render["blank_like_pages"] == []
    assert raw_render["page_sizes"] == [{"count": 539, "height": 2002, "width": 1547}]
    assert raw_render["edge_touch_pages"] == []
    assert template_render["pages"] == 380
    assert template_render["blank_like_pages"] == []
    assert template_render["page_sizes"] == [
        {"count": 380, "height": 2002, "width": 1547}
    ]
    assert template_render["edge_touch_pages"] == []
    for report in (raw_render, template_render):
        assert min(report["minimum_ink_margins_pixels"].values()) >= 150
    assert font_audit["passed"] is True
    assert font_audit["embedded_font_registrations"] == [
        "Noto Sans Symbols Regular",
        "Noto Sans Symbols Bold",
        "Roboto Mono Regular",
        "Roboto Mono Bold",
        "Roboto Mono Italic",
        "Roboto Mono BoldItalic",
    ]
    assert metrics["document_text_equality"] is True
    assert metrics["media_byte_equality"] is True
    assert metrics["media_files"] == EXPECTED_IMAGE_COUNT
    assert metrics["approximate_words"] >= 95_000
    assert visual_audit["docx"]["raw_media_matches_source"] is True
    assert visual_audit["docx"]["template_media_order_matches_raw"] is True
    assert visual_audit["docx"]["numbered_figure_caption_pairs"] == 25
    assert visual_audit["docx"]["alpha_images"] == 0
    assert visual_audit["pdf"]["images"] == EXPECTED_IMAGE_COUNT

    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            current_document = ET.fromstring(archive.read("word/document.xml"))
        tables = current_document.findall(f".//{{{WORD_NS}}}tbl")
        assert len(tables) == EXPECTED_TABLE_COUNT
        for table in tables:
            rows = table.findall(f"{{{WORD_NS}}}tr")
            assert rows
            header = rows[0].find(
                f"{{{WORD_NS}}}trPr/{{{WORD_NS}}}tblHeader"
            )
            assert header is not None
            assert header.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                "0",
                "false",
                "off",
            }
            for row in rows:
                keep_together = row.find(
                    f"{{{WORD_NS}}}trPr/{{{WORD_NS}}}cantSplit"
                )
                assert keep_together is not None
                assert keep_together.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                    "0",
                    "false",
                    "off",
                }

    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document_xml = archive.read("word/document.xml")
            font_table_xml = archive.read("word/fontTable.xml")

        document = ET.fromstring(document_xml)
        text = "".join(document.itertext())
        assert "Политика последовательности проверяет всю траекторию" in text
        assert text.count("approval -> resume -> execute -> audit") == 3
        assert "approval → resume → execute → audit" not in text
        assert b"Nova Mono" not in font_table_xml

        parent_by_child = {
            child: parent for parent in document.iter() for child in list(parent)
        }
        captions: list[tuple[int, ET.Element]] = []
        for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
            paragraph_value = "".join(
                node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
            ).strip()
            match = re.fullmatch(r"Рисунок (\d+)\. .+", paragraph_value)
            if match:
                captions.append((int(match.group(1)), paragraph))

        assert [number for number, _ in captions] == list(range(1, 26))
        for number, caption in captions:
            parent = parent_by_child[caption]
            siblings = list(parent)
            previous = next(
                (
                    candidate
                    for candidate in reversed(siblings[: siblings.index(caption)])
                    if candidate.tag == f"{{{WORD_NS}}}p"
                    and (
                        candidate.find(f".//{{{WORD_NS}}}drawing") is not None
                        or "".join(
                            node.text or ""
                            for node in candidate.findall(f".//{{{WORD_NS}}}t")
                        ).strip()
                    )
                ),
                None,
            )
            assert previous is not None, f"Figure {number} has no preceding paragraph"
            assert previous.find(f".//{{{WORD_NS}}}drawing") is not None
            image_keep_next = previous.find(
                f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext"
            )
            caption_keep_lines = caption.find(
                f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepLines"
            )
            assert image_keep_next is not None
            assert image_keep_next.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                "0",
                "false",
                "off",
            }
            assert caption_keep_lines is not None
            assert caption_keep_lines.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                "0",
                "false",
                "off",
            }


def test_template2000n_editorial_has_semantic_styles_and_image_alt_text() -> None:
    with ZipFile(EDITORIAL_TEMPLATE_DOCX) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        styles = ET.fromstring(archive.read("word/styles.xml"))
        relationships = ET.fromstring(archive.read("word/_rels/document.xml.rels"))

    style_ids = {
        node.attrib[f"{{{WORD_NS}}}styleId"] for node in styles.findall(f"{{{WORD_NS}}}style")
    }
    assert {
        "Title",
        "BodyText",
        "Heading1",
        "Heading2",
        "Heading3",
        "Heading4",
        "Style16",
        "Style17",
        "Style18",
        "Style20",
        "Style21",
        "Style23",
        "Style24",
        "Style28",
    } <= style_ids

    unstyled_non_empty = []
    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        text = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")).strip()
        style = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")
        if text and style is None:
            unstyled_non_empty.append(text)
    assert unstyled_non_empty == []

    image_properties = document.findall(f".//{{{DRAWING_NS}}}docPr")
    assert len(image_properties) == EXPECTED_IMAGE_COUNT
    assert all(node.attrib.get("descr", "").strip() for node in image_properties)

    hyperlinks = [
        node
        for node in relationships.findall(f"{{{PACKAGE_REL_NS}}}Relationship")
        if node.attrib.get("Type", "").endswith("/hyperlink")
    ]
    assert len(document.findall(f".//{{{WORD_NS}}}hyperlink")) >= 104
    assert len(hyperlinks) >= 104


def test_template2000n_prioritizes_list_style_over_inline_code_font() -> None:
    builder = ROOT / "docs/publisher/tools/build_template2000n_derivative.py"
    source = builder.read_text(encoding="utf-8")

    list_branch = source.index('elif paragraph.find("w:pPr/w:numPr", NS) is not None:')
    program_branch = source.index("elif paragraph_is_monospace(paragraph):")

    assert list_branch < program_branch


def test_template2000n_editorial_images_have_no_alpha_channel() -> None:
    alpha_images = []
    with ZipFile(EDITORIAL_TEMPLATE_DOCX) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/") or not name.lower().endswith(".png"):
                continue
            image = archive.read(name)
            color_type = image[25]
            if color_type in {4, 6} or b"tRNS" in image:
                alpha_images.append(name)

    assert alpha_images == []


def test_editorial_tables_repeat_headers_and_keep_rows_together() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        tables = document.findall(f".//{{{WORD_NS}}}tbl")
        assert tables
        for table in tables:
            rows = table.findall(f"{{{WORD_NS}}}tr")
            header = rows[0].find(f"{{{WORD_NS}}}trPr/{{{WORD_NS}}}tblHeader")
            assert header is not None
            assert header.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                "0",
                "false",
                "off",
            }
            for row in rows:
                keep_together = row.find(f"{{{WORD_NS}}}trPr/{{{WORD_NS}}}cantSplit")
                assert keep_together is not None
                assert keep_together.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                    "0",
                    "false",
                    "off",
                }


def test_editorial_docx_has_no_empty_numbered_paragraphs() -> None:
    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        empty_numbered = []
        for index, paragraph in enumerate(document.findall(f".//{{{WORD_NS}}}p")):
            text = "".join(
                node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
            ).strip()
            if (
                not text
                and paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}numPr") is not None
            ):
                empty_numbered.append(index)

        assert empty_numbered == []


def test_editorial_page_break_headings_have_no_leading_empty_paragraph() -> None:
    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraphs = document.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
        for index, paragraph in enumerate(paragraphs):
            if not active_on_off_property(paragraph, "pageBreakBefore"):
                continue
            assert index > 0
            previous_text = "".join(
                node.text or ""
                for node in paragraphs[index - 1].findall(f".//{{{WORD_NS}}}t")
            ).strip()
            assert previous_text


def test_editorial_code_blocks_avoid_orphaned_opening_lines() -> None:
    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraphs = document.findall(f".//{{{WORD_NS}}}body/{{{WORD_NS}}}p")
        code_blocks: list[list[ET.Element]] = []
        current_block: list[ET.Element] = []
        for paragraph in paragraphs:
            if paragraph_uses_monospace_font(paragraph):
                current_block.append(paragraph)
                continue
            if current_block:
                code_blocks.append(current_block)
                current_block = []
        if current_block:
            code_blocks.append(current_block)

        assert len(code_blocks) >= 30
        for block in code_blocks:
            assert all(active_on_off_property(paragraph, "keepLines") for paragraph in block)
            for index, paragraph in enumerate(block):
                expected = (
                    index < len(block) - 1
                    and (index + 1)
                    % sync_ru_docx_visuals.MAX_CODE_PARAGRAPHS_PER_KEEP_GROUP
                    != 0
                )
                assert active_on_off_property(paragraph, "keepNext") is expected


def test_inline_visual_titles_stay_with_their_images() -> None:
    for docx_path in (CURRENT_RAW_DOCX, CURRENT_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraphs = document.findall(f".//{{{WORD_NS}}}p")
        matched_titles = 0
        for index, paragraph in enumerate(paragraphs):
            image_properties = paragraph.find(f".//{{{DRAWING_NS}}}docPr")
            if image_properties is None:
                continue
            description = image_properties.attrib.get("descr", "").strip()
            previous_paragraph = next(
                (
                    candidate
                    for candidate in reversed(paragraphs[:index])
                    if "".join(
                        node.text or ""
                        for node in candidate.findall(f".//{{{WORD_NS}}}t")
                    ).strip()
                ),
                None,
            )
            if previous_paragraph is None:
                continue
            previous_text = "".join(
                node.text or ""
                for node in previous_paragraph.findall(f".//{{{WORD_NS}}}t")
            ).strip()
            if previous_text != description:
                continue

            matched_titles += 1
            keep_next = previous_paragraph.find(
                f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext"
            )
            assert keep_next is not None
            assert keep_next.attrib.get(f"{{{WORD_NS}}}val", "1") not in {
                "0",
                "false",
                "off",
            }

        assert matched_titles >= 10


def test_editorial_docx_has_print_navigation_language_and_metadata() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))
            settings = ET.fromstring(archive.read("word/settings.xml"))
            styles = ET.fromstring(archive.read("word/styles.xml"))
            core = ET.fromstring(archive.read("docProps/core.xml"))
            footer_names = [
                name for name in archive.namelist() if re.fullmatch(r"word/footer\d+\.xml", name)
            ]
            footer_xml = b"\n".join(archive.read(name) for name in footer_names)

        instructions = " ".join(
            node.text or "" for node in document.findall(f".//{{{WORD_NS}}}instrText")
        )
        document_text = "\n".join(
            node.text or "" for node in document.findall(f".//{{{WORD_NS}}}t")
        )
        style_languages = {
            node.attrib.get(f"{{{WORD_NS}}}val") for node in styles.findall(f".//{{{WORD_NS}}}lang")
        }

        assert 'TOC \\o "1-2"' in instructions
        assert settings.find(f"{{{WORD_NS}}}updateFields") is not None
        assert b"PAGE" in footer_xml
        assert "ru-RU" in style_languages
        assert core.findtext(f"{{{DC_NS}}}title") == "Архитектура безопасных ИИ-агентов"
        assert core.findtext(f"{{{DC_NS}}}language") == "ru-RU"
        assert core.findtext(f"{{{DC_NS}}}subject")
        assert core.findtext(f"{{{CP_NS}}}keywords")
        assert (
            len(re.findall(r"^Таблица \d+\. .+$", document_text, re.MULTILINE))
            == EXPECTED_TABLE_COUNT
        )


def test_editorial_heading_styles_define_pdf_outline_levels() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            styles = ET.fromstring(archive.read("word/styles.xml"))

        for level in range(1, 5):
            style = styles.find(f".//{{{WORD_NS}}}style[@{{{WORD_NS}}}styleId='Heading{level}']")
            assert style is not None
            outline_level = style.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}outlineLvl")
            assert outline_level is not None
            assert outline_level.attrib.get(f"{{{WORD_NS}}}val") == str(level - 1)


def test_editorial_toc_has_a_readable_static_result_before_word_updates_it() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraphs = document.findall(f".//{{{WORD_NS}}}p")
        paragraph_texts = [
            "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
            for paragraph in paragraphs
        ]
        text = "\n".join(paragraph_texts)
        first_body_heading = next(
            index
            for index, (paragraph, value) in enumerate(zip(paragraphs, paragraph_texts))
            if value == "Об авторе"
            and (style := paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")) is not None
            and style.attrib.get(f"{{{WORD_NS}}}val") == "Heading2"
        )
        toc_text = "\n".join(paragraph_texts[:first_body_heading])
        assert "Оглавление обновится при открытии документа." not in text
        assert "Часть I. От демо-агента к платформе" in toc_text
        assert "Глава 1. Почему агенту нужна платформа, а не магия" in toc_text
        assert "Глава 28. Проверочный список промышленного запуска" in toc_text
        assert "Приложение 5. Эталонный пакет и воспроизводимые упражнения" in toc_text


def test_template2000n_table_captions_use_caption_style_and_stay_with_tables() -> None:
    with ZipFile(EDITORIAL_TEMPLATE_DOCX) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    captions = []
    for paragraph in document.findall(f".//{{{WORD_NS}}}p"):
        value = "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")).strip()
        if re.fullmatch(r"Таблица \d+\. .+", value):
            captions.append(paragraph)

    assert len(captions) == EXPECTED_TABLE_COUNT
    for paragraph in captions:
        properties = paragraph.find(f"{{{WORD_NS}}}pPr")
        assert properties is not None
        style = properties.find(f"{{{WORD_NS}}}pStyle")
        assert style is not None
        assert style.attrib[f"{{{WORD_NS}}}val"] == "Style17"
        assert properties.find(f"{{{WORD_NS}}}keepNext") is not None
        assert properties.find(f"{{{WORD_NS}}}keepLines") is not None


def test_editorial_table_columns_have_readable_minimum_width() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        tables = document.findall(f".//{{{WORD_NS}}}tbl")
        assert len(tables) == EXPECTED_TABLE_COUNT
        for table_index, table in enumerate(tables, start=1):
            widths = [
                int(column.attrib[f"{{{WORD_NS}}}w"])
                for column in table.findall(f"{{{WORD_NS}}}tblGrid/{{{WORD_NS}}}gridCol")
            ]
            total_width = sum(widths)
            assert min(widths) / total_width >= 0.14, (
                f"table {table_index} in {docx_path.name} has a column narrower "
                "than 14% of the table width"
            )


def test_current_docx_exports_match_the_28_chapter_manuscript() -> None:
    for docx_path in (RAW_EDITORIAL_DOCX, EDITORIAL_TEMPLATE_DOCX):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraph_nodes = document.findall(f".//{{{WORD_NS}}}p")
        paragraphs = [
            "".join(node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t"))
            for paragraph in paragraph_nodes
        ]
        chapter_numbers = [
            int(match.group(1))
            for paragraph, paragraph_node in zip(paragraphs, paragraph_nodes)
            if (
                (style := paragraph_node.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")) is not None
                and style.attrib.get(f"{{{WORD_NS}}}val") == "Heading2"
                and (match := re.fullmatch(r"Глава (\d+)\. .+", paragraph.strip()))
            )
        ]
        text = "\n".join(paragraphs)

        assert chapter_numbers == list(range(1, 29))
        assert "create_support_ticket" not in text
        assert "Покрытие обязательным подтверждением и трассировкой" in text
        assert "except GatewayTimeout" in text
        figure_leads = re.findall(r"^На рисунке \d+ .+$", text, re.MULTILINE)
        assert len(figure_leads) == 25
        assert all("представлена схема" not in lead for lead in figure_leads)
        assert "дата доступа зафиксированы 15 июля 2026 года" in text
        assert "дата доступа зафиксированы 14 июля 2026 года" not in text
        assert text.count("После главы вы сможете:") == 28
        assert text.count("Артефакт главы:") == 28
        assert text.count("Предварительные условия.") == 8
        assert text.count("Отрицательная проверка.") == 8
        assert len(re.findall(r"Листинг \d+\.", text)) >= 30
        assert "agentic_internal_risk:" in text
        assert "control_plane_readiness:" in text


def test_editorial_page_breaks_only_start_reader_facing_sections() -> None:
    for docx_path in (FINAL_LAYOUT_DOCX,):
        with ZipFile(docx_path) as archive:
            document = ET.fromstring(archive.read("word/document.xml"))

        paragraphs = document.findall(f".//{{{WORD_NS}}}p")
        page_breaks = []
        previous_text = None
        for paragraph in paragraphs:
            text = "".join(
                node.text or "" for node in paragraph.findall(f".//{{{WORD_NS}}}t")
            ).strip()
            page_break = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pageBreakBefore")
            if page_break is not None:
                value = page_break.attrib.get(f"{{{WORD_NS}}}val", "true")
                if value not in {"false", "0", "off"}:
                    style = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")
                    page_breaks.append((text, style.attrib.get(f"{{{WORD_NS}}}val")))
                    assert text
                    assert previous_text
            previous_text = text

        assert len(page_breaks) == 38
        assert all(style in {"Heading1", "Heading2"} for _, style in page_breaks)
        chapter_breaks = {
            int(match.group(1))
            for text, _ in page_breaks
            if (match := re.fullmatch(r"Глава (\d+)\. .+", text))
        }
        assert chapter_breaks == set(range(1, 29))


def test_baseline_docx_artifacts_preserve_the_same_manuscript_image_order() -> None:
    manuscript = EDITORIAL_MANUSCRIPT.read_text(encoding="utf-8")
    relative_paths = re.findall(r"^!\[[^\]]+\]\((visuals/[^)]+)\)$", manuscript, re.MULTILINE)
    raw_targets, _ = ordered_embedded_images(CURRENT_RAW_DOCX)
    template_targets, _ = ordered_embedded_images(CURRENT_TEMPLATE_DOCX)

    assert len(relative_paths) == EXPECTED_IMAGE_COUNT
    assert len(raw_targets) == EXPECTED_IMAGE_COUNT
    assert template_targets == raw_targets


def test_numbered_figure_captions_follow_images_and_stay_with_them() -> None:
    with ZipFile(EDITORIAL_TEMPLATE_DOCX) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))

    paragraphs = document.findall(f".//{{{WORD_NS}}}p")
    numbered_pairs = 0
    for index, paragraph in enumerate(paragraphs):
        if paragraph.find(f".//{{{DRAWING_NS}}}docPr") is None:
            continue
        next_paragraph = next(
            (
                candidate
                for candidate in paragraphs[index + 1 :]
                if "".join(
                    node.text or "" for node in candidate.findall(f".//{{{WORD_NS}}}t")
                ).strip()
            ),
            None,
        )
        assert next_paragraph is not None
        next_text = "".join(
            node.text or "" for node in next_paragraph.findall(f".//{{{WORD_NS}}}t")
        ).strip()
        if re.fullmatch(r"Рисунок \d+\. .+", next_text) is None:
            continue

        numbered_pairs += 1
        image_style = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")
        caption_style = next_paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")
        assert image_style is not None
        assert image_style.attrib[f"{{{WORD_NS}}}val"] == "Style28"
        assert paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepNext") is not None
        assert caption_style is not None
        assert caption_style.attrib[f"{{{WORD_NS}}}val"] == "Style17"
        assert next_paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}keepLines") is not None

    assert numbered_pairs == 25
